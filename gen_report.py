# -*- coding: utf-8 -*-
"""
Generate Ho Chi Minh City University of Post and Telecommunications
CNPM Course Project Report - Bus Management System
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION

# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────

def set_page_setup(section):
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.header_distance = Cm(1.27)
    section.footer_distance = Cm(1.27)

def add_heading(doc, text, level=1, bold=True, size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def add_paragraph(doc, text="", bold=False, size=13, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=3, space_after=6, indent=0):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before    = Pt(space_before)
    pf.space_after    = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing   = 1.08
    if indent:
        pf.left_indent = Cm(indent)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
    return p

def add_bullet(doc, text, size=13, level=0):
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.75 + level * 0.75)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def add_code(doc, text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(1)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(size)
    return p

def add_table_bordered(doc, rows, cols, data=None):
    t = doc.add_table(rows=rows, cols=cols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if data:
        for r in range(rows):
            for c in range(cols):
                cell = t.cell(r, c)
                cell.text = data[r][c]
                cell.paragraphs[0].runs[0].font.size = Pt(11)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return t

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def page_break(doc):
    doc.add_page_break()

# ─────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────

def build_cover(doc):
    # Institution lines - 16pt bold centered
    for line in ["BỘ KHOA HỌC VÀ CÔNG NGHỆ", "HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(line)
        r.bold = True
        r.font.size = Pt(16)

    # Separator line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(12)
    r = p.add_run("─" * 50)
    r.font.size = Pt(14)

    # Title "BÁO CÁO ĐỒ ÁN MÔN HỌC" - 36pt bold
    for word in ["BÁO CÁO", "ĐỒ ÁN MÔN HỌC"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(word)
        r.bold = True
        r.font.size = Pt(36)

    # Subject line - 16pt
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run("MÔN HỌC: CÔNG NGHỆ PHẦN MỀM")
    r.bold = True
    r.font.size = Pt(16)

    # Topic line - 16pt
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run("ĐỀ TÀI: HỆ THỐNG QUẢN LÝ & ĐẶT VÉ XE KHÁCH LIÊN TỈNH")
    r.bold = True
    r.font.size = Pt(16)

    # Spacer
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)

    # GVHD
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(0)
    r1 = p.add_run("Giảng viên hướng dẫn: ")
    r1.bold = True
    r1.font.size = Pt(13)
    r2 = p.add_run("Nguyễn Thị Bích Nguyên")
    r2.font.size = Pt(13)

    # SV label
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run("Thực hiện bởi nhóm sinh viên, bao gồm:")
    r.bold = True
    r.font.size = Pt(13)

    # Student rows
    students = [
        ("Ngô Đức Tú",       "N23DCAT073", ""),
        ("Lê Vũ Hảo",        "N23DCCN154", "D23CQCN03-N"),
        ("Đặng Phương Nam",  "N23DCCN039", ""),
    ]
    for name, mssv, lop in students:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        # Tab-separated: name + tab + MSSV + (optional class)
        content = f"{name:<22}\t{mssv:<12}"
        if lop:
            content += f"\t{lop}"
        r = p.add_run(content)
        r.font.size = Pt(13)

    # Spacer
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)

    # City and date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("TP.HCM, tháng 06 / 2026")
    r.bold = True
    r.font.size = Pt(12)

    page_break(doc)

# ─────────────────────────────────────────────
# TABLE OF CONTENTS
# ─────────────────────────────────────────────

def build_toc(doc):
    add_heading(doc, "MỤC LỤC", level=1, size=16, space_before=6, space_after=12)
    toc_entries = [
        ("MỤC LỤC", "3"),
        ("DANH SÁCH HÌNH, BẢNG", "3"),
        ("DANH MỤC TỪ VIẾT TẮT", "3"),
        ("CHƯƠNG I. TỔNG QUAN", "4"),
        ("    I. Giới thiệu đề tài", "4"),
        ("    II. Cơ sở lý thuyết", "4"),
        ("CHƯƠNG II. PHÂN TÍCH NỘI DUNG, YÊU CẦU", "5"),
        ("    I. Giới thiệu quy trình đặt vé", "5"),
        ("    II. Giới thiệu quy trình quản lý", "6"),
        ("    III. Yêu cầu chức năng nghiệp vụ", "7"),
        ("    IV. Yêu cầu chức năng hệ thống và yêu cầu chất lượng", "10"),
        ("CHƯƠNG III. PHÂN TÍCH THIẾT KẾ", "11"),
        ("    I. Sơ đồ use case", "11"),
        ("    II. Sơ đồ hoạt động", "13"),
        ("    III. Thiết kế cơ sở dữ liệu", "14"),
        ("    IV. Thiết kế giao diện", "17"),
        ("    V. Thiết kế xử lý", "18"),
        ("CHƯƠNG IV. PHÁT TRIỂN / THỰC THI", "20"),
        ("    I. Màn hình đăng nhập / đăng ký", "20"),
        ("    II. Màn hình tìm chuyến & đặt vé", "21"),
        ("    III. Màn hình thanh toán", "22"),
        ("    IV. Màn hình Admin Dashboard", "23"),
        ("    V. Màn hình quản lý chuyến xe & phân công", "25"),
        ("CHƯƠNG V. TRIỂN KHAI", "26"),
        ("    I. Cài đặt", "26"),
        ("    II. Thử nghiệm", "26"),
        ("CHƯƠNG VI. KẾT LUẬN", "28"),
        ("    I. Kết quả đã thực hiện", "28"),
        ("    II. Ưu khuyết điểm", "28"),
        ("    III. Hướng mở rộng trong tương lai", "29"),
        ("TÀI LIỆU THAM KHẢO", "29"),
    ]
    for entry, page in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
        r = p.add_run(entry + "\t" + page)
        r.font.size = Pt(12)

    page_break(doc)

# ─────────────────────────────────────────────
# FIGURES & TABLES LIST
# ─────────────────────────────────────────────

def build_figures_tables(doc):
    add_heading(doc, "DANH SÁCH HÌNH, BẢNG", level=1, size=16, space_before=6, space_after=12)

    add_heading(doc, "Danh sách hình", level=2, size=14, space_before=6, space_after=6)
    figures = [
        ("Hình 1.1", "Kiến trúc tổng quan hệ thống Bus Management"),
        ("Hình 1.2", "Kiến trúc Spring Boot MVC"),
        ("Hình 3.1", "Sơ đồ Use Case tổng thể"),
        ("Hình 3.2", "Sơ đồ Use Case - Khách hàng"),
        ("Hình 3.3", "Sơ đồ Use Case - Quản trị viên"),
        ("Hình 3.4", "Sơ đồ hoạt động - Quy trình đặt vé"),
        ("Hình 3.5", "Sơ đồ hoạt động - Thanh toán VNPay"),
        ("Hình 3.6", "Sơ đồ hoạt động - Quản lý chuyến xe"),
        ("Hình 3.7", "Mô hình ERD hệ thống"),
        ("Hình 3.8", "Sơ đồ quan hệ bảng (Database Diagram)"),
        ("Hình 3.9", "Giao diện trang đăng nhập"),
        ("Hình 3.10", "Giao diện trang đăng ký"),
        ("Hình 3.11", "Giao diện tìm chuyến xe"),
        ("Hình 3.12", "Giao diện chọn ghế ngồi"),
        ("Hình 3.13", "Giao diện thanh toán VNPay"),
        ("Hình 3.14", "Giao diện Admin Dashboard"),
        ("Hình 3.15", "Giao diện quản lý người dùng"),
        ("Hình 3.16", "Giao diện quản lý xe"),
        ("Hình 3.17", "Giao diện quản lý tuyến đường"),
        ("Hình 3.18", "Giao diện quản lý chuyến xe"),
        ("Hình 3.19", "Giao diện phân công tài xế"),
        ("Hình 4.1", "Màn hình đăng nhập / đăng ký"),
        ("Hình 4.2", "Màn hình tìm chuyến"),
        ("Hình 4.3", "Màn hình chọn điểm đón / trả"),
        ("Hình 4.4", "Màn hình chọn ghế"),
        ("Hình 4.5", "Màn hình xác nhận đặt vé"),
        ("Hình 4.6", "Màn hình thanh toán VNPay"),
        ("Hình 4.7", "Màn hình Admin Dashboard"),
        ("Hình 4.8", "Màn hình quản lý chuyến xe"),
        ("Hình 4.9", "Màn hình phân công tài xế & phụ xe"),
    ]
    for stt, desc in figures:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(f"{stt}: {desc}")
        r.font.size = Pt(12)

    add_heading(doc, "Danh sách bảng", level=2, size=14, space_before=12, space_after=6)
    tables_list = [
        ("Bảng 2.1", "Yêu cầu chức năng - Khách hàng"),
        ("Bảng 2.2", "Yêu cầu chức năng - Quản trị viên"),
        ("Bảng 2.3", "Yêu cầu chức năng hệ thống"),
        ("Bảng 3.1", "Danh sách tác nhân (Actors)"),
        ("Bảng 3.2", "Danh sách Use Case"),
        ("Bảng 3.3", "Bảng User"),
        ("Bảng 3.4", "Bảng Passenger"),
        ("Bảng 3.5", "Bảng Bus"),
        ("Bảng 3.6", "Bảng Seat"),
        ("Bảng 3.7", "Bảng Route"),
        ("Bảng 3.8", "Bảng Trip"),
        ("Bảng 3.9", "Bảng Ticket"),
        ("Bảng 3.10", "Bảng Payment"),
        ("Bảng 3.11", "Bảng Employee"),
        ("Bảng 3.12", "Bảng TripAssignment"),
        ("Bảng 3.13", "Bảng Maintenance"),
        ("Bảng 3.14", "Bảng Cargo"),
        ("Bảng 3.15", "Bảng AuditLog"),
        ("Bảng 5.1", "Bảng cài đặt chức năng"),
        ("Bảng 5.2", "Tài khoản test"),
    ]
    for stt, desc in tables_list:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(f"{stt}: {desc}")
        r.font.size = Pt(12)

    page_break(doc)

# ─────────────────────────────────────────────
# ABBREVIATIONS
# ─────────────────────────────────────────────

def build_abbreviations(doc):
    add_heading(doc, "DANH MỤC TỪ VIẾT TẮT", level=1, size=16, space_before=6, space_after=12)
    abbrevs = [
        ("API",    "Application Programming Interface - Giao diện lập trình ứng dụng"),
        ("CNPM",   "Công Nghệ Phần Mềm - Software Engineering"),
        ("CORS",   "Cross-Origin Resource Sharing"),
        ("CRUD",   "Create, Read, Update, Delete"),
        ("CSS",    "Cascading Style Sheets"),
        ("DTO",    "Data Transfer Object - Đối tượng truyền dữ liệu"),
        ("ERD",    "Entity-Relationship Diagram - Sơ đồ thực thể liên kết"),
        ("HTML",   "HyperText Markup Language"),
        ("HTTP",   "HyperText Transfer Protocol"),
        ("HTTPS",  "HyperText Transfer Protocol Secure"),
        ("JPA",    "Java Persistence API"),
        ("JSON",   "JavaScript Object Notation"),
        ("JWT",    "JSON Web Token"),
        ("MySQL",  "My Structured Query Language"),
        ("ORM",    "Object-Relational Mapping"),
        ("REST",   "Representational State Transfer"),
        ("SSE",    "Server-Sent Events"),
        ("SQL",    "Structured Query Language"),
        ("SSO",    "Single Sign-On"),
        ("UI",     "User Interface - Giao diện người dùng"),
        ("UX",     "User Experience - Trải nghiệm người dùng"),
        ("VNPay",  "Vietnam Payment Gateway - Cổng thanh toán VNPay"),
        ("Vite",   "Next Generation Frontend Tooling"),
    ]
    for abbr, meaning in abbrevs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(4.5))
        r1 = p.add_run(f"{abbr}\t")
        r1.bold = True
        r1.font.size = Pt(12)
        r2 = p.add_run(meaning)
        r2.font.size = Pt(12)

    page_break(doc)

# ─────────────────────────────────────────────
# CHAPTER I - TỔNG QUAN
# ─────────────────────────────────────────────

def build_chapter1(doc):
    add_heading(doc, "CHƯƠNG I. TỔNG QUAN", level=1, size=16, space_before=6, space_after=12)

    # I. Giới thiệu đề tài
    add_heading(doc, "I. Giới thiệu đề tài", level=2, size=14, space_before=8, space_after=6)

    add_heading(doc, "1. Mục tiêu của đề tài", level=3, size=13, bold=True, space_before=6, space_after=4)
    add_paragraph(doc,
        "Xây dựng Hệ thống quản lý & đặt vé xe khách liên tỉnh (Bus Management System) nhằm "
        "số hóa quy trình đặt vé, quản lý chuyến xe và thanh toán trực tuyến, giúp hành khách "
        "có thể tìm chuyến xe, chọn ghế và thanh toán tiện lợi từ bất kỳ thiết bị nào có Internet. "
        "Hệ thống cũng hỗ trợ bộ phận quản lý giám sát hoạt động kinh doanh vận tải theo thời gian thực."
    )
    add_paragraph(doc, "Cụ thể, hệ thống hướng đến các mục tiêu:", bold=False, space_before=3)
    bullets_1 = [
        "Cho phép hành khách tra cứu và đặt vé xe khách liên tỉnh 24/7.",
        "Tích hợp cổng thanh toán trực tuyến VNPay hỗ trợ nhiều phương thức (ATM, Visa, QR).",
        "Hỗ trợ thanh toán COD (Cash On Delivery) cho hành khách chưa có tài khoản ngân hàng.",
        "Quản lý chuyến xe, phân công tài xế và phụ xe, quản lý bảo trì xe.",
        "Dashboard quản trị thời gian thực với thông báo SSE khi có đặt vé mới.",
        "Đảm bảo tính an toàn thông tin qua xác thực JWT và mã hóa dữ liệu thanh toán.",
    ]
    for b in bullets_1:
        add_bullet(doc, b, size=13)

    add_heading(doc, "2. Phạm vi áp dụng", level=3, size=13, bold=True, space_before=8, space_after=4)
    add_paragraph(doc,
        "Hệ thống được triển khai phục vụ các đối tượng sau:"
    )
    bullets_2 = [
        "Hành khách có nhu cầu di chuyển liên tỉnh bằng xe khách.",
        "Nhân viên quản lý (Admin) phụ trách vận hành, giám sát và điều phối.",
        "Nhân viên trực tiếp điều hành (Dispatcher) theo dõi chuyến xe.",
    ]
    for b in bullets_2:
        add_bullet(doc, b, size=13)

    add_heading(doc, "3. Nền tảng kỹ thuật", level=3, size=13, bold=True, space_before=8, space_after=4)
    add_paragraph(doc,
        "Hệ thống được phát triển theo kiến trúc Full-Stack với hai module chính: Backend (Spring Boot) "
        "và Frontend (React + Vite + TypeScript). Cơ sở dữ liệu sử dụng MySQL 8.0. "
        "Chi tiết công nghệ được trình bày trong Bảng 1.1 dưới đây."
    )
    add_paragraph(doc, "Bảng 1.1: Bảng công nghệ sử dụng", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=4)

    tech_data = [
        ["Thành phần",   "Công nghệ",                       "Phiên bản",  "Mục đích sử dụng"],
        ["Ngôn ngữ BE",  "Java",                             "17+",        "Backend logic, REST API"],
        ["Framework BE", "Spring Boot",                      "3.2.0",      "Web framework chính"],
        ["ORM",          "Spring Data JPA / Hibernate",      "6.x",        "Ánh xạ đối tượng - quan hệ"],
        ["Database",     "MySQL",                            "8.0",        "Lưu trữ dữ liệu"],
        ["Auth",         "Spring Security + JWT",            "JJWT 0.11.5","Xác thực & phân quyền"],
        ["Payment",      "VNPay Sandbox API",                "v2.1.0",     "Cổng thanh toán trực tuyến"],
        ["Build BE",     "Maven",                            "3.9+",       "Quản lý dependency"],
        ["Ngôn ngữ FE",  "TypeScript",                       "5.6",        "Frontend type-safe"],
        ["Framework FE", "React",                            "18.3",       "Giao diện người dùng"],
        ["Build FE",     "Vite",                             "5.4",        "Bundler & dev server"],
        ["Styling",      "Tailwind CSS",                     "3.4",        "Utility-first CSS"],
        ["State",        "Zustand",                          "4.5",        "Quản lý trạng thái"],
        ["Data Fetch",   "Axios + React Query",              "5.x",        "Gọi API + caching"],
        ["Charts",       "Recharts",                         "2.8",        "Biểu đồ dashboard"],
        ["Icons",        "Lucide React",                     "0.400+",     "Biểu tượng UI"],
    ]
    t = add_table_bordered(doc, len(tech_data), 4, tech_data)
    for i, cell in enumerate(t.rows[0].cells):
        shade_cell(cell, "D9E1F2")
        cell.paragraphs[0].runs[0].bold = True

    # II. Cơ sở lý thuyết
    add_heading(doc, "II. Cơ sở lý thuyết", level=2, size=14, space_before=12, space_after=6)

    add_heading(doc, "1. Kiến trúc Spring Boot MVC", level=3, size=13, bold=True, space_before=6, space_after=4)
    add_paragraph(doc,
        "Spring Boot là framework phổ biến nhất trong hệ sinh thái Java Enterprise, cho phép xây dựng "
        "REST API theo mô hình MVC (Model - View - Controller). Trong hệ thống này, Controller tiếp nhận "
        "request HTTP từ client, chuyển đến Service xử lý nghiệp vụ, Service gọi Repository để truy vấn "
        "database thông qua JPA/Hibernate, và trả kết quả về cho client dưới dạng JSON."
    )
    add_paragraph(doc,
        "Lợi ích của Spring Boot: tự động cấu hình (auto-configuration), nhúng server (embedded Tomcat), "
        "hỗ trợ dependency injection, kiểm thử dễ dàng, và hệ sinh thái phong phú với Spring Security, "
        "Spring Data JPA, Spring Validation."
    )

    add_heading(doc, "2. JWT (JSON Web Token)", level=3, size=13, bold=True, space_before=8, space_after=4)
    add_paragraph(doc,
        "JWT là chuẩn mở (RFC 7519) cho phép truyền thông tin an toàn giữa hai bên dưới dạng JSON. "
        "JWT gồm 3 phần: Header (thuật toán), Payload (dữ liệu claims), Signature (chữ ký HMAC-SHA256). "
        "Trong hệ thống, mỗi khi người dùng đăng nhập thành công, server tạo JWT chứa username và role, "
        "gửi về cho client. Client lưu JWT trong localStorage và gửi kèm trong header "
        "'Authorization: Bearer <token>' cho mọi request tiếp theo. JwtAuthenticationFilter của Spring "
        "Security kiểm tra token trước khi cho phép truy cập resource."
    )

    add_heading(doc, "3. React & Vite", level=3, size=13, bold=True, space_before=8, space_after=4)
    add_paragraph(doc,
        "React là thư viện giao diện phổ biến nhất hiện nay, sử dụng component-based architecture và "
        "virtual DOM để tối ưu hiệu năng render. Vite là công cụ build thế hệ mới với HMR (Hot Module "
        "Replacement) cực nhanh, sử dụng native ES Module trong dev mode. TypeScript mang lại type-safety "
        "giúp giảm lỗi runtime và hỗ trợ IntelliSense mạnh mẽ."
    )

    add_heading(doc, "4. VNPay Payment Gateway", level=3, size=13, bold=True, space_before=8, space_after=4)
    add_paragraph(doc,
        "VNPay là cổng thanh toán trực tuyến phổ biến tại Việt Nam, hỗ trợ thanh toán qua thẻ ATM nội địa, "
        "thẻ quốc tế Visa/MasterCard, và mã QR ngân hàng. Quy trình tích hợp gồm 3 bước: "
        "(1) Server tạo URL thanh toán với các tham số mã hóa HMAC-SHA512; "
        "(2) Khách hàng được chuyển sang cổng VNPay và nhập thông tin thanh toán; "
        "(3) VNPay redirect về Return URL với trạng thái giao dịch và gửi IPN callback server-to-server để "
        "cập nhật trạng thái vé."
    )

    add_heading(doc, "5. Server-Sent Events (SSE)", level=3, size=13, bold=True, space_before=8, space_after=4)
    add_paragraph(doc,
        "SSE là công nghệ cho phép server gửi cập nhật đến client qua kết nối HTTP một chiều (server → client). "
        "Khác với WebSocket (hai chiều), SSE đơn giản hơn và phù hợp cho trường hợp server cần push thông báo "
        "đến client mà không cần client gửi dữ liệu lên server. Trong hệ thống, AdminNotificationService sử "
        "dụng SseEmitter của Spring để quản lý các kết nối SSE từ các tab Admin Dashboard và push "
        "thông báo real-time khi có sự kiện."
    )

    add_heading(doc, "6. JPA / Hibernate ORM", level=3, size=13, bold=True, space_before=8, space_after=4)
    add_paragraph(doc,
        "Java Persistence API (JPA) là specification cho Object-Relational Mapping (ORM) trong Java. "
        "Hibernate là implementation phổ biến nhất của JPA. JPA giúp ánh xạ các class Java (Entity) thành "
        "các bảng trong MySQL, tự động tạo câu lệnh SQL, hỗ trợ relationship mapping (OneToMany, ManyToOne, "
        "ManyToMany), cascading và lazy loading. Spring Data JPA mở rộng thêm Repository pattern, cung cấp "
        "các method như findBy..., save, delete mà không cần viết SQL thủ công."
    )

    page_break(doc)

# ─────────────────────────────────────────────
# CHAPTER II - PHÂN TÍCH NỘI DUNG, YÊU CẦU
# ─────────────────────────────────────────────

def build_chapter2(doc):
    add_heading(doc, "CHƯƠNG II. PHÂN TÍCH NỘI DUNG, YÊU CẦU", level=1, size=16, space_before=6, space_after=12)

    # I. Giới thiệu quy trình 1 - Đặt vé
    add_heading(doc, "I. Giới thiệu quy trình đặt vé", level=2, size=14, space_before=8, space_after=6)
    add_paragraph(doc,
        "Quy trình đặt vé là luồng chính của hệ thống, bắt đầu từ khi hành khách tra cứu chuyến xe và kết "
        "thúc khi vé được thanh toán và xác nhận. Hệ thống hỗ trợ hai phương thức thanh toán: "
        "(1) Thanh toán trực tuyến qua VNPay; (2) Thanh toán COD (trả tiền mặt khi lên xe). "
        "Quy trình gồm 6 bước chính:"
    )

    steps = [
        ("Bước 1: Tìm kiếm chuyến xe",
         "Hành khách nhập điểm đi, điểm đến và ngày khởi hành. Hệ thống truy vấn bảng Route và Trip, "
         "lọc các chuyến có trạng thái SCHEDULED và ngày khớp, trả về danh sách kèm thông tin xe, "
         "giờ khởi hành, giá vé."),
        ("Bước 2: Chọn điểm đón",
         "Hành khách chọn thành phố điểm đón, sau đó chọn 1 trong danh sách điểm đón cụ thể. "
         "Dữ liệu điểm đón/trả được khởi tạo sẵn từ DataInitializer, gồm 80+ điểm phủ 15 thành phố."),
        ("Bước 3: Chọn điểm trả",
         "Tương tự bước 2, hành khách chọn thành phố và điểm trả cụ thể."),
        ("Bước 4: Chọn ghế ngồi",
         "Hệ thống hiển thị sơ đồ ghế dạng grid 5 cột. Các ghế đã đặt hiển thị màu đỏ, ghế trống màu xanh, "
         "ghế đang chọn màu xanh dương. Khi hành khách chọn ghế, hệ thống gọi API kiểm tra ghế có "
         "còn trống không (dùng SELECT FOR UPDATE để chống race condition) và tạm lock ghế trong 15 phút."),
        ("Bước 5: Xác nhận thông tin & chọn phương thức thanh toán",
         "Hành khách xem lại tóm tắt vé: thông tin hành khách, chuyến xe, ghế, điểm đón/trả, giá. "
         "Chọn thanh toán VNPay (trực tuyến) hoặc COD (trả tiền mặt khi lên xe). "
         "Hệ thống tạo Ticket với trạng thái BOOKED (VNPay) hoặc CONFIRMED (COD)."),
        ("Bước 6: Thanh toán & xác nhận",
         "Nếu chọn VNPay: hệ thống tạo URL thanh toán, hành khách được chuyển sang cổng VNPay. "
         "Sau khi thanh toán, VNPay redirect về Return URL. Backend xác thực HMAC-SHA512, cập nhật "
         "Payment và Ticket sang trạng thái PAID, gửi SSE notification đến Admin Dashboard. "
         "Nếu chọn COD: vé được xác nhận ngay, Admin nhận SSE notification để gọi điện xác nhận với khách."),
    ]
    for title, desc in steps:
        add_heading(doc, title, level=3, size=13, bold=True, space_before=6, space_after=3)
        add_paragraph(doc, desc, space_before=2)

    # II. Quy trình quản lý
    add_heading(doc, "II. Giới thiệu quy trình quản lý", level=2, size=14, space_before=12, space_after=6)
    add_paragraph(doc,
        "Hệ thống cung cấp bộ công cụ quản lý toàn diện cho Quản trị viên (Admin) để vận hành hoạt động "
        "kinh doanh vận tải. Các quy trình quản lý chính bao gồm:"
    )

    mgmt = [
        ("Quản lý tuyến đường",
         "Admin có thể tạo, cập nhật, xóa tuyến đường. Mỗi tuyến gồm: điểm đi, điểm đến, khoảng cách (km), "
         "thời gian ước tính (phút), giá vé cơ bản. Hệ thống hiện có 4 tuyến: Hà Nội↔TP.HCM (1700km, 500k), "
         "Hà Nội↔Đà Nẵng (800km, 300k), TP.HCM↔Đà Nẵng (900km, 350k), Hà Nội↔Hải Phòng (100km, 80k)."),
        ("Quản lý xe khách",
         "Admin quản lý danh sách 20 xe bus với các loại: Limousine (28 chỗ), Sleeper cao cấp (34 chỗ), "
         "Sleeper thường (40 chỗ), Seat cao cấp (45 chỗ), Seat thường (50 chỗ). Mỗi xe có biển số, loại xe, "
         "số ghế, ngày đăng kiểm, ngày hết hạn bảo hiểm. Hệ thống tự động cảnh báo khi bảo hiểm sắp hết hạn "
         "(trong vòng 30 ngày)."),
        ("Quản lý chuyến xe",
         "Admin tạo chuyến xe gồm: chọn tuyến, chọn xe, nhập giờ khởi hành. Hệ thống kiểm tra xem xe đó "
         "có chuyến đang chạy trùng lịch không (overlap detection). Trạng thái chuyến: SCHEDULED → RUNNING "
         "→ COMPLETED, hoặc DELAYED, CANCELLED."),
        ("Quản lý nhân sự",
         "Admin quản lý 20 nhân viên gồm: 10 tài xế (5 VIP kinh nghiệm 11-15 năm, 5 thường 3-7 năm) "
         "và 10 phụ xe (kinh nghiệm 1-6 năm). Mỗi nhân viên có: họ tên, SĐT, quê quán, loại, năm kinh nghiệm."),
        ("Phân công tài xế & phụ xe",
         "Admin phân công tài xế và phụ xe cho từng chuyến. Mỗi chuyến cần ít nhất 1 tài xế, có thể "
         "thêm 1 phụ xe. Hệ thống kiểm tra tài xế không bị phân công chuyến khác trùng giờ."),
        ("Quản lý vé",
         "Admin xem danh sách tất cả vé, lọc theo trạng thái (BOOKED, HOLD, PAID, CONFIRMED, CANCELLED, "
         "REFUNDED, EXPIRED), xác nhận vé COD, hủy vé, xem chi tiết vé kèm thông tin hành khách, "
         "thông tin thanh toán, ghế đã đặt."),
        ("Quản lý bảo trì xe",
         "Admin ghi nhận lịch sử bảo trì: ngày bảo trì, mô tả công việc, chi phí, trạng thái "
         "(SCHEDULED, IN_PROGRESS, COMPLETED)."),
        ("Quản lý hàng hóa",
         "Admin quản lý vận chuyển hàng hóa: thông tin người gửi, người nhận, loại hàng, trọng lượng, "
         "phí vận chuyển, trạng thái (PENDING, IN_TRANSIT, DELIVERED, CANCELLED)."),
        ("Nhật ký hệ thống (Audit Log)",
         "Mọi thao tác thêm / sửa / xóa của Admin đều được ghi lại trong bảng AuditLog: ai làm, "
         "bảng nào, dữ liệu cũ, dữ liệu mới, thời gian. Giúp kiểm soát và truy vết hoạt động."),
    ]
    for title, desc in mgmt:
        add_heading(doc, title, level=3, size=13, bold=True, space_before=6, space_after=3)
        add_paragraph(doc, desc, space_before=2)

    # III. Yêu cầu chức năng nghiệp vụ
    add_heading(doc, "III. Yêu cầu chức năng nghiệp vụ", level=2, size=14, space_before=12, space_after=6)

    add_heading(doc, "1. Yêu cầu chức năng - Đối tượng: Khách hàng", level=3, size=13, bold=True, space_before=6, space_after=6)
    add_paragraph(doc, "Bảng 2.1: Yêu cầu chức năng nghiệp vụ - Khách hàng", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=3, space_after=4)

    cust_data = [
        ["STT", "Công việc", "Loại", "Quy định / Mô tả", "Ghi chú"],
        ["1", "Đăng nhập", "Tra cứu",
         "Tài khoản phải tồn tại trong DB, bắt buộc nhập email và mật khẩu. "
         "Sai mật khẩu 5 lần → khóa tài khoản 30 phút.",
         "JWT token 1 giờ"],
        ["2", "Đăng ký tài khoản", "Lưu trữ",
         "Bắt buộc: email, họ tên, SĐT, mật khẩu. Email & SĐT phải duy nhất. "
         "Mật khẩu mã hóa BCrypt trước khi lưu.",
         "Auto tạo Passenger record"],
        ["3", "Tra cứu chuyến xe", "Tra cứu",
         "Lọc theo điểm đi, điểm đến, ngày. Chỉ hiển thị chuyến SCHEDULED. "
         "Sắp xếp theo giờ khởi hành tăng dần.",
         ""],
        ["4", "Xem chi tiết chuyến", "Tra cứu",
         "Hiển thị: tuyến, xe, giờ khởi hành/đến, giá vé, sơ đồ ghế (màu: xanh=trống, "
         "đỏ=đã đặt, xanh dương=đang chọn).",
         ""],
        ["5", "Đặt vé", "Lưu trữ",
         "Chọn ghế trống → tạo Ticket trạng thái BOOKED/HOLD. Lock ghế trong 15 phút. "
         "Ghế đã đặt không thể đặt lại (SELECT FOR UPDATE).",
         "Race-condition-safe"],
        ["6", "Thanh toán VNPay", "Lưu trữ",
         "Tạo URL thanh toán VNPay với HMAC-SHA512. IPN callback → cập nhật Payment=PAID, "
         "Ticket=PAID. Return URL → hiển thị kết quả.",
         "Idempotent"],
        ["7", "Thanh toán COD", "Lưu trữ",
         "Tạo Ticket CONFIRMED ngay. Admin nhận SSE notification để gọi xác nhận với khách.",
         "Real-time notification"],
        ["8", "Xem danh sách vé của tôi", "Tra cứu",
         "Hiển thị tất cả vé của khách hàng đang đăng nhập, kèm trạng thái, thông tin chuyến, "
         "phương thức thanh toán.",
         ""],
        ["9", "Hủy vé", "Cập nhật",
         "Chỉ vé ở trạng thái BOOKED/HOLD/CONFIRMED mới được hủy. "
         "Không được hủy trong vòng 15 phút trước giờ khởi hành.",
         ""],
        ["10", "Xem / cập nhật hồ sơ", "Cập nhật",
         "Xem thông tin cá nhân: email, họ tên, SĐT. Cập nhật họ tên và SĐT. "
         "Không cho phép đổi email.",
         ""],
    ]
    t = add_table_bordered(doc, len(cust_data), 5, cust_data)
    for cell in t.rows[0].cells:
        shade_cell(cell, "D9E1F2")
        cell.paragraphs[0].runs[0].bold = True

    add_heading(doc, "2. Yêu cầu chức năng - Đối tượng: Quản trị viên", level=3, size=13, bold=True, space_before=10, space_after=6)
    add_paragraph(doc, "Bảng 2.2: Yêu cầu chức năng nghiệp vụ - Quản trị viên", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=3, space_after=4)

    admin_data = [
        ["STT", "Công việc", "Loại", "Quy định / Mô tả", "Ghi chú"],
        ["1", "Đăng nhập", "Tra cứu",
         "Tài khoản role=ADMIN, xác thực JWT. Dashboard hiển thị sau khi đăng nhập.",
         "JWT token 1 giờ"],
        ["2", "Xem Dashboard", "Tra cứu",
         "Hiển thị: tổng Users/Buses/Routes, chuyến hôm nay, biểu đồ phân bổ, "
         "cảnh báo bảo hiểm sắp hết hạn.",
         "Real-time SSE"],
        ["3", "Quản lý Users", "CRUD",
         "Thêm / sửa / khóa / mở khóa / reset mật khẩu người dùng. "
         "Tài khoản bị khóa: status=LOCKED.",
         ""],
        ["4", "Quản lý Buses", "CRUD",
         "Thêm / sửa / xóa xe. Cập nhật trạng thái: ACTIVE, MAINTENANCE, RETIRED. "
         "Auto tạo bản ghi Seat khi thêm xe.",
         "Auto-generate seats"],
        ["5", "Quản lý Routes", "CRUD",
         "Thêm / sửa / xóa tuyến. Tính giá tự động: basePrice = distanceKm × 300₫.",
         ""],
        ["6", "Quản lý Trips", "CRUD",
         "Thêm / sửa / xóa chuyến. Kiểm tra overlap: xe không được phép có 2 chuyến "
         "trùng thời gian.",
         "Overlap detection"],
        ["7", "Phân công nhân sự", "Lưu trữ",
         "Phân công tài xế & phụ xe cho chuyến. Kiểm tra tài xế không bị trùng lịch. "
         "Mỗi chuyến cần ít nhất 1 tài xế.",
         ""],
        ["8", "Quản lý nhân viên", "CRUD",
         "Thêm / sửa / xóa nhân viên. Lọc theo loại: DRIVER, ASSISTANT, TECHNICIAN, "
         "DISPATCHER, MANAGER.",
         ""],
        ["9", "Quản lý vé", "Tra cứu",
         "Xem danh sách tất cả vé, lọc theo trạng thái. Xác nhận vé COD. Hủy vé. "
         "Xem chi tiết: hành khách, thanh toán, ghế.",
         ""],
        ["10", "Quản lý Maintenance", "CRUD",
         "Ghi nhận lịch bảo trì: ngày, mô tả, chi phí, trạng thái.",
         ""],
        ["11", "Quản lý Cargo", "CRUD",
         "Quản lý vận chuyển hàng hóa: PENDING → IN_TRANSIT → DELIVERED / CANCELLED.",
         ""],
        ["12", "Nhận thông báo real-time", "Thông báo",
         "SSE kết nối từ Admin Dashboard. Popup toast khi có booking mới (BOOKING_CREATED) "
         "hoặc thanh toán VNPay thành công (PAYMENT_VNPAY_SUCCESS).",
         "SSE real-time"],
        ["13", "Xem nhật ký hệ thống", "Tra cứu",
         "Xem AuditLog: hành động, bảng, dữ liệu cũ/mới, thời gian. Lọc theo ngày, user.",
         ""],
    ]
    t = add_table_bordered(doc, len(admin_data), 5, admin_data)
    for cell in t.rows[0].cells:
        shade_cell(cell, "D9E1F2")
        cell.paragraphs[0].runs[0].bold = True

    # IV. Yêu cầu hệ thống
    add_heading(doc, "IV. Yêu cầu chức năng hệ thống và yêu cầu chất lượng", level=2, size=14, space_before=12, space_after=6)

    add_heading(doc, "1. Yêu cầu chức năng hệ thống", level=3, size=13, bold=True, space_before=6, space_after=6)
    add_paragraph(doc, "Bảng 2.3: Yêu cầu chức năng hệ thống", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=3, space_after=4)

    sys_data = [
        ["STT", "Yêu cầu", "Mức ưu tiên", "Mô tả chi tiết"],
        ["1", "Xác thực & Phân quyền", "Bắt buộc",
         "JWT Bearer token. Phân quyền: ADMIN truy cập /api/admin/**, CUSTOMER truy cập /api/private/**. "
         "Token expire sau 1 giờ."],
        ["2", "Bảo mật mật khẩu", "Bắt buộc",
         "BCrypt hashing với strength=10. Không lưu plain text."],
        ["3", "CORS", "Bắt buộc",
         "Cho phép frontend dev (localhost:5173) và production domain truy cập API."],
        ["4", "Validation", "Bắt buộc",
         "Tất cả request DTO phải có @Valid annotation. Email đúng format, SĐT 10-11 số."],
        ["5", "Error handling", "Bắt buộc",
         "GlobalExceptionHandler trả về ApiErrorResponse{status, message, timestamp}."],
        ["6", "Logging", "Cao",
         "Ghi log startup, shutdown, lỗi hệ thống vào backend.log."],
        ["7", "SSE broadcaster", "Cao",
         "CopyOnWriteArrayList để quản lý SseEmitter thread-safe. Auto cleanup khi client disconnect."],
        ["8", "Payment idempotent", "Bắt buộc",
         "Trước khi cập nhật Payment→SUCCESS, kiểm tra chưa có bản ghi SUCCESS nào với txnRef đó."],
        ["9", "Concurrency control", "Bắt buộc",
         "Dùng @Lock(PESSIMISTIC_WRITE) - SELECT FOR UPDATE khi đặt ghế để chống race condition."],
        ["10", "Responsive UI", "Cao",
         "Giao diện web thích ứng desktop/mobile. Tailwind CSS breakpoints: sm, md, lg."],
        ["11", "Performance", "Trung bình",
         "Thời gian phản hồi API < 500ms với database có 10,000 bản ghi."],
        ["12", "Availability", "Cao",
         "Hệ thống hoạt động 24/7. Backend restart không mất dữ liệu (MySQL persistent)."],
    ]
    t = add_table_bordered(doc, len(sys_data), 4, sys_data)
    for cell in t.rows[0].cells:
        shade_cell(cell, "D9E1F2")
        cell.paragraphs[0].runs[0].bold = True

    add_heading(doc, "2. Yêu cầu chất lượng", level=3, size=13, bold=True, space_before=10, space_after=4)
    qual = [
        "Tính đúng đắn: mọi giao dịch đặt vé phải được xử lý chính xác, không trùng ghế, không mất vé.",
        "Tính an toàn: thông tin thanh toán VNPay phải được xác thực HMAC-SHA512, không chấp nhận request giả mạo.",
        "Tính khả dụng: hệ thống phải phản hồi nhanh, giao diện thân thiện, dễ sử dụng.",
        "Tính bảo trì: code có comment rõ ràng, chia layer Controller → Service → Repository rõ ràng.",
        "Tính mở rộng: kiến trúc modular cho phép thêm tính năng mới (SMS notification, mobile app) dễ dàng.",
    ]
    for q in qual:
        add_bullet(doc, q, size=13)

    page_break(doc)

# ─────────────────────────────────────────────
# CHAPTER III - PHÂN TÍCH THIẾT KẾ
# ─────────────────────────────────────────────

def build_chapter3(doc):
    add_heading(doc, "CHƯƠNG III. PHÂN TÍCH THIẾT KẾ", level=1, size=16, space_before=6, space_after=12)

    # I. Sơ đồ Use Case
    add_heading(doc, "I. Sơ đồ use case", level=2, size=14, space_before=8, space_after=6)

    add_heading(doc, "1. Danh sách tác nhân (Actors)", level=3, size=13, bold=True, space_before=6, space_after=6)
    add_paragraph(doc, "Bảng 3.1: Danh sách tác nhân", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=3, space_after=4)
    actors = [
        ["Mã tác nhân", "Tên tác nhân", "Mô tả"],
        ["AC1", "Khách chưa đăng nhập", "Hành khách chưa có tài khoản, chỉ xem được danh sách chuyến"],
        ["AC2", "Khách hàng (CUSTOMER)", "Hành khách đã đăng ký, có tài khoản để đặt vé và thanh toán"],
        ["AC3", "Quản trị viên (ADMIN)", "Người quản lý hệ thống, có quyền CRUD tất cả entities"],
        ["AC4", "Nhân viên điều phối (DISPATCHER)", "Theo dõi và giám sát chuyến xe trong ca trực"],
        ["AC5", "Hệ thống VNPay", "Cổng thanh toán bên thứ ba, xử lý giao dịch tài chính"],
    ]
    t = add_table_bordered(doc, len(actors), 3, actors)
    for cell in t.rows[0].cells:
        shade_cell(cell, "D9E1F2")
        cell.paragraphs[0].runs[0].bold = True

    add_heading(doc, "2. Danh sách Use Case", level=3, size=13, bold=True, space_before=10, space_after=6)
    add_paragraph(doc, "Bảng 3.2: Danh sách Use Case", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=3, space_after=4)
    usecases = [
        ["Mã UC", "Tên Use Case", "Tác nhân chính", "Mô tả"],
        ["UC01", "Đăng nhập", "Khách chưa đăng nhập / Khách hàng / Admin",
         "Xác thực tài khoản, trả về JWT token"],
        ["UC02", "Đăng ký tài khoản", "Khách chưa đăng nhập",
         "Tạo tài khoản mới, auto tạo Passenger record"],
        ["UC03", "Tra cứu chuyến xe", "Khách chưa đăng nhập / Khách hàng",
         "Lọc chuyến theo điểm đi, điểm đến, ngày"],
        ["UC04", "Xem chi tiết chuyến & sơ đồ ghế", "Khách hàng",
         "Hiển thị thông tin chuyến và trạng thái từng ghế"],
        ["UC05", "Đặt vé", "Khách hàng",
         "Chọn ghế, tạo Ticket, lock ghế 15 phút"],
        ["UC06", "Thanh toán VNPay", "Khách hàng + VNPay",
         "Tạo URL thanh toán, xử lý IPN, cập nhật trạng thái vé"],
        ["UC07", "Thanh toán COD", "Khách hàng",
         "Xác nhận vé ngay, gửi notification cho Admin"],
        ["UC08", "Xem vé của tôi", "Khách hàng",
         "Xem danh sách vé đã đặt và trạng thái"],
        ["UC09", "Hủy vé", "Khách hàng",
         "Hủy vé BOOKED/HOLD/CONFIRMED không trong vòng 15 phút trước giờ khởi hành"],
        ["UC10", "Xem / cập nhật hồ sơ", "Khách hàng",
         "Xem và cập nhật họ tên, SĐT cá nhân"],
        ["UC11", "Xem Dashboard", "Admin",
         "Xem thống kê tổng quan, biểu đồ, cảnh báo"],
        ["UC12", "Quản lý người dùng", "Admin",
         "CRUD users, khóa/mở khóa, reset mật khẩu"],
        ["UC13", "Quản lý xe khách", "Admin",
         "CRUD buses, cập nhật trạng thái xe"],
        ["UC14", "Quản lý tuyến đường", "Admin",
         "CRUD routes, tự động tính giá"],
        ["UC15", "Quản lý chuyến xe", "Admin",
         "CRUD trips, kiểm tra overlap"],
        ["UC16", "Phân công tài xế & phụ xe", "Admin",
         "Gán nhân viên cho chuyến, kiểm tra trùng lịch"],
        ["UC17", "Quản lý nhân viên", "Admin",
         "CRUD employees, lọc theo loại"],
        ["UC18", "Quản lý vé", "Admin",
         "Xem tất cả vé, xác nhận COD, hủy vé"],
        ["UC19", "Quản lý bảo trì xe", "Admin",
         "CRUD maintenance records"],
        ["UC20", "Quản lý hàng hóa", "Admin",
         "CRUD cargo, cập nhật trạng thái vận chuyển"],
        ["UC21", "Nhận thông báo real-time", "Admin",
         "SSE kết nối, popup toast khi có sự kiện"],
        ["UC22", "Xem nhật ký hệ thống", "Admin",
         "Xem AuditLog, lọc theo user và thời gian"],
    ]
    t = add_table_bordered(doc, len(usecases), 4, usecases)
    for cell in t.rows[0].cells:
        shade_cell(cell, "D9E1F2")
        cell.paragraphs[0].runs[0].bold = True

    add_heading(doc, "3. Mô tả Use Case chi tiết", level=3, size=13, bold=True, space_before=10, space_after=6)

    # UC05 Đặt vé
    add_heading(doc, "UC05 - Đặt vé", level=3, size=13, bold=True, space_before=4, space_after=3)
    add_paragraph(doc, "Tác nhân: Khách hàng")
    add_paragraph(doc, "Tiền điều kiện: Khách hàng đã đăng nhập, đang xem sơ đồ ghế của một chuyến xe có trạng thái SCHEDULED.")
    add_paragraph(doc, "Luồng chính:")
    uc05_steps = [
        "1. Khách hàng nhấn chọn một ghế trống (màu xanh).",
        "2. Hệ thống gửi POST /api/private/tickets với tripId, seatId, passengerName, passengerPhone, pickupPointId, dropoffPointId.",
        "3. Backend: JwtService xác thực token → extract userId.",
        "4. Backend: TicketService gọi findByIdForUpdate(seatId) để lock dòng trong DB.",
        "5. Backend: Kiểm tra seat chưa có Ticket trạng thái active (BOOKED/HOLD/PAID/CONFIRMED).",
        "6. Backend: Kiểm tra trip có trạng thái SCHEDULED và departureTime > now + 15 phút.",
        "7. Backend: Tạo Passenger (nếu chưa có), tạo Ticket(status=BOOKED), save.",
        "8. Backend: Trả về TicketResponse cho client.",
        "9. Client hiển thị ghế đã chọn màu xanh dương, hiển thị form xác nhận.",
    ]
    for s in uc05_steps:
        add_bullet(doc, s, size=12)
    add_paragraph(doc, "Luồng phụ:", bold=True)
    add_bullet(doc, "Nếu ghế đã được đặt (race condition): trả lỗi 409 Conflict.", size=12)
    add_bullet(doc, "Nếu trip sắp khởi hành (<15 phút): trả lỗi 400 Bad Request.", size=12)
    add_paragraph(doc, "Hậu điều kiện: Ticket tạo thành công, ghế chuyển sang trạng thái BOOKED, vé có 15 phút để thanh toán.")

    # UC06 VNPay
    add_heading(doc, "UC06 - Thanh toán VNPay", level=3, size=13, bold=True, space_before=6, space_after=3)
    add_paragraph(doc, "Tác nhân: Khách hàng, Hệ thống VNPay")
    add_paragraph(doc, "Tiền điều kiện: Ticket đang ở trạng thái BOOKED/HOLD.")
    add_paragraph(doc, "Luồng chính:")
    uc06_steps = [
        "1. Khách hàng chọn thanh toán VNPay, nhấn 'Thanh toán ngay'.",
        "2. Frontend gửi POST /api/private/payment/vnpay/create với ticketId.",
        "3. Backend: TicketService tìm Ticket, kiểm tra trạng thái, tính expiryTime = now + 15 phút.",
        "4. Backend: VnpayService tạo map tham số: vnp_Amount, vnp_Command, vnp_CreateDate, vnp_IpAddr, vnp_Locale, vnp_OrderInfo, vnp_TxnRef (format: TICKET-{ticketId}-{timestamp}), vnp_Version, vnp_ReturnUrl.",
        "5. Backend: Mã hóa toàn bộ query string bằng HMAC-SHA512 với secret key, gán vnp_SecureHash.",
        "6. Backend: Tạo Payment record (status=PENDING), trả URL thanh toán VNPay về client.",
        "7. Frontend redirect sang URL VNPay.",
        "8. Khách hàng nhập thông tin thẻ tại cổng VNPay.",
        "9. VNPay xử lý giao dịch, redirect về Return URL: GET /api/public/payment/vnpay/return.",
        "10. Backend: VnpayService verify vnp_SecureHash, kiểm tra ResponseCode.",
        "11. Backend: Nếu ResponseCode='00' (thành công), gọi IPN handler xử lý: kiểm tra idempotent → cập nhật Payment→SUCCESS, Ticket→PAID.",
        "12. Backend: Gửi SSE event 'PAYMENT_VNPAY_SUCCESS' đến Admin Dashboard.",
        "13. Frontend hiển thị trang kết quả: thành công hoặc thất bại kèm mã vé.",
    ]
    for s in uc06_steps:
        add_bullet(doc, s, size=12)
    add_paragraph(doc, "Hậu điều kiện: Nếu thanh toán thành công: Ticket=PAID, Payment=SUCCESS. Nếu thất bại: Ticket giữ nguyên BOOKED, khách có thể thử lại.")

    # II. Sơ đồ hoạt động
    add_heading(doc, "II. Sơ đồ hoạt động", level=2, size=14, space_before=12, space_after=6)

    add_heading(doc, "1. Sơ đồ hoạt động - Quy trình đặt vé", level=3, size=13, bold=True, space_before=6, space_after=4)
    add_paragraph(doc,
        "Sơ đồ hoạt động (Activity Diagram) mô tả luồng xử lý từ khi khách hàng tìm chuyến xe đến khi "
        "nhận được vé đã thanh toán. Luồng bắt đầu từ khách hàng đăng nhập (hoặc tiếp tục với vai trò "
        "khách chưa đăng nhập cho bước tra cứu), nhập điểm đi/đến/ngày để tìm chuyến. "
        "Sau khi chọn chuyến, khách chọn điểm đón và điểm trả từ danh sách 80+ điểm. "
        "Hệ thống hiển thị sơ đồ ghế, khách chọn ghế trống. "
        "Hệ thống kiểm tra ghế (SELECT FOR UPDATE), nếu ghế còn trống → tạo Ticket BOOKED. "
        "Khách chọn phương thức thanh toán: VNPay hoặc COD. "
        "Nếu VNPay: tạo URL thanh toán → redirect VNPay → xác thực → cập nhật PAID → gửi SSE notification. "
        "Nếu COD: xác nhận CONFIRMED → gửi SSE notification cho Admin. "
        "Cuối cùng, hệ thống hiển thị thông tin vé hoàn chỉnh cho khách."
    )

    add_heading(doc, "2. Sơ đồ hoạt động - Quy trình thanh toán VNPay", level=3, size=13, bold=True, space_before=8, space_after=4)
    add_paragraph(doc,
        "Sơ đồ mô tả chi tiết tương tác giữa Client, Backend và VNPay. "
        "Khi khách hàng chọn thanh toán VNPay, Backend tạo request với HMAC-SHA512 signature, "
        "gửi URL thanh toán về client. Client redirect sang cổng VNPay. "
        "Sau khi thanh toán, VNPay gửi đồng thời: (1) redirect về Return URL (phía client) và "
        "(2) IPN callback server-to-server. Backend xử lý IPN trước (idempotent check) để cập nhật "
        "trạng thái Payment và Ticket. Return URL chỉ hiển thị kết quả cho khách. "
        "Mọi trạng thái đều được ghi log, đảm bảo tính truy vết."
    )

    add_heading(doc, "3. Sơ đồ hoạt động - Quản lý chuyến xe", level=3, size=13, bold=True, space_before=8, space_after=4)
    add_paragraph(doc,
        "Admin đăng nhập → vào Dashboard → chọn 'Quản lý chuyến xe'. "
        "Để tạo chuyến mới, Admin chọn tuyến (điểm đi - điểm đến), chọn xe, nhập giờ khởi hành. "
        "Backend kiểm tra overlap: với xe đã chọn, kiểm tra tất cả chuyến SCHEDULED/RUNNING, "
        "nếu có khoảng thời gian chồng lấn → báo lỗi. "
        "Nếu không trùng → tạo Trip(status=SCHEDULED). "
        "Sau đó Admin phân công tài xế và phụ xe: chọn chuyến, gắn nhân viên. "
        "Hệ thống kiểm tra nhân viên không có assignment trùng giờ → tạo TripAssignment. "
        "Khi đến giờ khởi hành, Admin cập nhật trạng thái chuyến sang RUNNING. "
        "Khi hoàn thành → COMPLETED. Admin có thể hủy chuyến → CANCELLED."
    )

    # III. Thiết kế CSDL
    add_heading(doc, "III. Thiết kế cơ sở dữ liệu", level=2, size=14, space_before=12, space_after=6)

    add_heading(doc, "1. Mô hình ERD", level=3, size=13, bold=True, space_before=6, space_after=4)
    add_paragraph(doc,
        "Mô hình ERD (Entity-Relationship Diagram) mô tả các thực thể và mối quan hệ giữa chúng. "
        "Hệ thống có 14 thực thể chính được mô tả dưới đây, liên kết qua các quan hệ 1-1, 1-N, N-1, N-N."
    )

    add_heading(doc, "2. Mô tả các bảng (Table Structure)", level=3, size=13, bold=True, space_before=8, space_after=6)
    add_paragraph(doc, "Bảng 3.3: Bảng User", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=3, space_after=4)
    user_cols = [
        ["STT", "Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
        ["1", "id", "BIGINT", "PK, AUTO_INCREMENT", "Khóa chính"],
        ["2", "username", "VARCHAR(50)", "UNIQUE, NOT NULL", "Tên đăng nhập"],
        ["3", "password_hash", "VARCHAR(255)", "NOT NULL", "BCrypt hash của mật khẩu"],
        ["4", "email", "VARCHAR(100)", "UNIQUE, NOT NULL", "Email người dùng"],
        ["5", "phone", "VARCHAR(15)", "NULL", "Số điện thoại"],
        ["6", "status", "ENUM", "NOT NULL, DEFAULT 'ACTIVE'", "ACTIVE / INACTIVE / LOCKED"],
        ["7", "created_at", "TIMESTAMP", "NOT NULL", "Ngày tạo tài khoản"],
        ["8", "updated_at", "TIMESTAMP", "NULL", "Ngày cập nhật gần nhất"],
    ]
    t = add_table_bordered(doc, len(user_cols), 5, user_cols)
    for cell in t.rows[0].cells:
        shade_cell(cell, "D9E1F2")
        cell.paragraphs[0].runs[0].bold = True

    add_paragraph(doc, "Bảng 3.4: Bảng Passenger", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)
    pass_cols = [
        ["STT", "Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
        ["1", "id", "BIGINT", "PK, AUTO_INCREMENT", "Khóa chính"],
        ["2", "user_id", "BIGINT", "FK → User(id), UNIQUE", "Liên kết 1-1 với User"],
        ["3", "full_name", "VARCHAR(100)", "NOT NULL", "Họ tên đầy đủ"],
        ["4", "phone", "VARCHAR(15)", "NOT NULL", "SĐT hành khách"],
        ["5", "email", "VARCHAR(100)", "NULL", "Email"],
        ["6", "id_card", "VARCHAR(20)", "NULL", "CMND/CCCD"],
    ]
    t = add_table_bordered(doc, len(pass_cols), 5, pass_cols)
    for cell in t.rows[0].cells:
        shade_cell(cell, "D9E1F2")
        cell.paragraphs[0].runs[0].bold = True

    add_paragraph(doc, "Bảng 3.5: Bảng Bus", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)
    bus_cols = [
        ["STT", "Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
        ["1", "id", "BIGINT", "PK, AUTO_INCREMENT", "Khóa chính"],
        ["2", "license_plate", "VARCHAR(20)", "UNIQUE, NOT NULL", "Biển số xe"],
        ["3", "bus_type", "ENUM", "NOT NULL", "LIMOUSINE / SLEEPER / SEAT"],
        ["4", "total_seats", "INT", "NOT NULL", "Tổng số ghế"],
        ["5", "status", "ENUM", "NOT NULL, DEFAULT 'ACTIVE'", "ACTIVE / MAINTENANCE / RETIRED"],
        ["6", "registration_expiry", "DATE", "NULL", "Ngày hết hạn đăng kiểm"],
        ["7", "insurance_expiry", "DATE", "NULL", "Ngày hết hạn bảo hiểm"],
    ]
    t = add_table_bordered(doc, len(bus_cols), 5, bus_cols)
    for cell in t.rows[0].cells:
        shade_cell(cell, "D9E1F2")
        cell.paragraphs[0].runs[0].bold = True

    add_paragraph(doc, "Bảng 3.6: Bảng Seat", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)
    seat_cols = [
        ["STT", "Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
        ["1", "id", "BIGINT", "PK, AUTO_INCREMENT", "Khóa chính"],
        ["2", "bus_id", "BIGINT", "FK → Bus(id), NOT NULL", "Xe chứa ghế này"],
        ["3", "seat_number", "VARCHAR(5)", "NOT NULL", "Số ghế (1A, 1B...)"],
        ["4", "position_x", "INT", "NOT NULL", "Vị trí cột trong grid (1-5)"],
        ["5", "position_y", "INT", "NOT NULL", "Vị trí hàng trong grid"],
        ["6", "seat_type", "ENUM", "NOT NULL", "STANDARD / VIP"],
    ]
    t = add_table_bordered(doc, len(seat_cols), 5, seat_cols)
    for cell in t.rows[0].cells:
        shade_cell(cell, "D9E1F2")
        cell.paragraphs[0].runs[0].bold = True

    add_paragraph(doc, "Bảng 3.7: Bảng Route", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)
    route_cols = [
        ["STT", "Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
        ["1", "id", "BIGINT", "PK, AUTO_INCREMENT", "Khóa chính"],
        ["2", "origin", "VARCHAR(100)", "NOT NULL", "Điểm đi"],
        ["3", "destination", "VARCHAR(100)", "NOT NULL", "Điểm đến"],
        ["4", "distance_km", "DOUBLE", "NOT NULL", "Khoảng cách (km)"],
        ["5", "estimated_duration_min", "INT", "NOT NULL", "Thời gian ước tính (phút)"],
        ["6", "base_price", "DECIMAL(10,2)", "NOT NULL", "Giá vé cơ bản (VNĐ)"],
    ]
    t = add_table_bordered(doc, len(route_cols), 5, route_cols)
    for cell in t.rows[0].cells:
        shade_cell(cell, "D9E1F2")
        cell.paragraphs[0].runs[0].bold = True

    add_paragraph(doc, "Bảng 3.8: Bảng Trip", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)
    trip_cols = [
        ["STT", "Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
        ["1", "id", "BIGINT", "PK, AUTO_INCREMENT", "Khóa chính"],
        ["2", "route_id", "BIGINT", "FK → Route(id), NOT NULL", "Tuyến đường"],
        ["3", "bus_id", "BIGINT", "FK → Bus(id), NOT NULL", "Xe khách"],
        ["4", "departure_time", "TIMESTAMP", "NOT NULL", "Giờ khởi hành"],
        ["5", "arrival_time", "TIMESTAMP", "NOT NULL", "Giờ đến (tự tính)"],
        ["6", "status", "ENUM", "NOT NULL", "SCHEDULED/RUNNING/COMPLETED/DELAYED/CANCELLED"],
        ["7", "current_price", "DECIMAL(10,2)", "NOT NULL", "Giá vé hiện tại"],
    ]
    t = add_table_bordered(doc, len(trip_cols), 5, trip_cols)
    for cell in t.rows[0].cells:
        shade_cell(cell, "D9E1F2")
        cell.paragraphs[0].runs[0].bold = True

    add_paragraph(doc, "Bảng 3.9: Bảng Ticket", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)
    ticket_cols = [
        ["STT", "Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
        ["1", "id", "BIGINT", "PK, AUTO_INCREMENT", "Khóa chính"],
        ["2", "trip_id", "BIGINT", "FK → Trip(id), NOT NULL", "Chuyến xe"],
        ["3", "seat_id", "BIGINT", "FK → Seat(id), NOT NULL", "Ghế ngồi"],
        ["4", "passenger_id", "BIGINT", "FK → Passenger(id), NOT NULL", "Hành khách"],
        ["5", "price", "DECIMAL(10,2)", "NOT NULL", "Giá vé đã đặt"],
        ["6", "status", "ENUM", "NOT NULL", "BOOKED/HOLD/PAID/CONFIRMED/CANCELLED/EXPIRED/REFUNDED"],
        ["7", "pickup_point", "VARCHAR(200)", "NULL", "Điểm đón"],
        ["8", "dropoff_point", "VARCHAR(200)", "NULL", "Điểm trả"],
        ["9", "booking_code", "VARCHAR(20)", "UNIQUE", "Mã đặt vé: BUS-YYYYMMDD-XXXXX"],
        ["10", "booking_expires_at", "TIMESTAMP", "NULL", "Thời hạn thanh toán"],
        ["11", "created_at", "TIMESTAMP", "NOT NULL", "Ngày đặt"],
    ]
    t = add_table_bordered(doc, len(ticket_cols), 5, ticket_cols)
    for cell in t.rows[0].cells:
        shade_cell(cell, "D9E1F2")
        cell.paragraphs[0].runs[0].bold = True

    add_paragraph(doc, "Bảng 3.10: Bảng Payment", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)
    pay_cols = [
        ["STT", "Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
        ["1", "id", "BIGINT", "PK, AUTO_INCREMENT", "Khóa chính"],
        ["2", "ticket_id", "BIGINT", "FK → Ticket(id), NOT NULL", "Vé liên kết"],
        ["3", "amount", "DECIMAL(10,2)", "NOT NULL", "Số tiền thanh toán"],
        ["4", "payment_method", "ENUM", "NOT NULL", "CASH/CARD/MOMO/BANK/VNPAY"],
        ["5", "status", "ENUM", "NOT NULL", "PENDING/SUCCESS/FAILED/REFUNDED"],
        ["6", "vnp_txn_ref", "VARCHAR(100)", "NULL, UNIQUE", "Mã giao dịch VNPay"],
        ["7", "vnp_bank_code", "VARCHAR(20)", "NULL", "Mã ngân hàng thanh toán"],
        ["8", "vnp_response_code", "VARCHAR(10)", "NULL", "Mã phản hồi VNPay"],
        ["9", "created_at", "TIMESTAMP", "NOT NULL", "Thời điểm tạo"],
        ["10", "updated_at", "TIMESTAMP", "NULL", "Thời điểm cập nhật"],
    ]
    t = add_table_bordered(doc, len(pay_cols), 5, pay_cols)
    for cell in t.rows[0].cells:
        shade_cell(cell, "D9E1F2")
        cell.paragraphs[0].runs[0].bold = True

    add_paragraph(doc, "Bảng 3.11: Bảng Employee", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)
    emp_cols = [
        ["STT", "Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
        ["1", "id", "BIGINT", "PK, AUTO_INCREMENT", "Khóa chính"],
        ["2", "full_name", "VARCHAR(100)", "NOT NULL", "Họ tên nhân viên"],
        ["3", "phone", "VARCHAR(15)", "NOT NULL", "SĐT"],
        ["4", "hometown", "VARCHAR(100)", "NULL", "Quê quán"],
        ["5", "experience_years", "INT", "NOT NULL, DEFAULT 0", "Số năm kinh nghiệm"],
        ["6", "employee_type", "ENUM", "NOT NULL", "DRIVER/ASSISTANT/TECHNICIAN/DISPATCHER/MANAGER"],
        ["7", "status", "ENUM", "NOT NULL, DEFAULT 'ACTIVE'", "ACTIVE / INACTIVE"],
    ]
    t = add_table_bordered(doc, len(emp_cols), 5, emp_cols)
    for cell in t.rows[0].cells:
        shade_cell(cell, "D9E1F2")
        cell.paragraphs[0].runs[0].bold = True

    add_paragraph(doc, "Bảng 3.12: Bảng TripAssignment", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)
    assign_cols = [
        ["STT", "Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
        ["1", "id", "BIGINT", "PK, AUTO_INCREMENT", "Khóa chính"],
        ["2", "trip_id", "BIGINT", "FK → Trip(id), NOT NULL", "Chuyến xe"],
        ["3", "employee_id", "BIGINT", "FK → Employee(id), NOT NULL", "Nhân viên"],
        ["4", "assignment_role", "ENUM", "NOT NULL", "DRIVER / ASSISTANT"],
        ["5", "assigned_at", "TIMESTAMP", "NOT NULL", "Thời điểm phân công"],
    ]
    t = add_table_bordered(doc, len(assign_cols), 5, assign_cols)
    for cell in t.rows[0].cells:
        shade_cell(cell, "D9E1F2")
        cell.paragraphs[0].runs[0].bold = True

    add_paragraph(doc, "Bảng 3.13: Bảng Maintenance", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)
    maint_cols = [
        ["STT", "Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
        ["1", "id", "BIGINT", "PK, AUTO_INCREMENT", "Khóa chính"],
        ["2", "bus_id", "BIGINT", "FK → Bus(id), NOT NULL", "Xe được bảo trì"],
        ["3", "description", "TEXT", "NULL", "Mô tả công việc"],
        ["4", "cost", "DECIMAL(10,2)", "NULL", "Chi phí bảo trì"],
        ["5", "maintenance_date", "DATE", "NOT NULL", "Ngày bảo trì"],
        ["6", "status", "ENUM", "NOT NULL", "SCHEDULED / IN_PROGRESS / COMPLETED"],
    ]
    t = add_table_bordered(doc, len(maint_cols), 5, maint_cols)
    for cell in t.rows[0].cells:
        shade_cell(cell, "D9E1F2")
        cell.paragraphs[0].runs[0].bold = True

    add_paragraph(doc, "Bảng 3.14: Bảng Cargo", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)
    cargo_cols = [
        ["STT", "Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
        ["1", "id", "BIGINT", "PK, AUTO_INCREMENT", "Khóa chính"],
        ["2", "trip_id", "BIGINT", "FK → Trip(id), NULL", "Chuyến vận chuyển"],
        ["3", "sender_name", "VARCHAR(100)", "NOT NULL", "Tên người gửi"],
        ["4", "sender_phone", "VARCHAR(15)", "NOT NULL", "SĐT người gửi"],
        ["5", "receiver_name", "VARCHAR(100)", "NOT NULL", "Tên người nhận"],
        ["6", "receiver_phone", "VARCHAR(15)", "NOT NULL", "SĐT người nhận"],
        ["7", "cargo_type", "VARCHAR(100)", "NOT NULL", "Loại hàng hóa"],
        ["8", "weight", "DOUBLE", "NULL", "Trọng lượng (kg)"],
        ["9", "fee", "DECIMAL(10,2)", "NOT NULL", "Phí vận chuyển"],
        ["10", "status", "ENUM", "NOT NULL", "PENDING / IN_TRANSIT / DELIVERED / CANCELLED"],
    ]
    t = add_table_bordered(doc, len(cargo_cols), 5, cargo_cols)
    for cell in t.rows[0].cells:
        shade_cell(cell, "D9E1F2")
        cell.paragraphs[0].runs[0].bold = True

    add_paragraph(doc, "Bảng 3.15: Bảng AuditLog", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)
    audit_cols = [
        ["STT", "Tên cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
        ["1", "id", "BIGINT", "PK, AUTO_INCREMENT", "Khóa chính"],
        ["2", "user_id", "BIGINT", "FK → User(id), NULL", "Người thực hiện"],
        ["3", "action", "VARCHAR(50)", "NOT NULL", "CREATE / UPDATE / DELETE"],
        ["4", "table_name", "VARCHAR(100)", "NOT NULL", "Tên bảng bị thay đổi"],
        ["5", "record_id", "BIGINT", "NULL", "ID bản ghi bị thay đổi"],
        ["6", "old_values", "TEXT", "NULL", "JSON dữ liệu cũ"],
        ["7", "new_values", "TEXT", "NULL", "JSON dữ liệu mới"],
        ["8", "created_at", "TIMESTAMP", "NOT NULL", "Thời điểm thực hiện"],
    ]
    t = add_table_bordered(doc, len(audit_cols), 5, audit_cols)
    for cell in t.rows[0].cells:
        shade_cell(cell, "D9E1F2")
        cell.paragraphs[0].runs[0].bold = True

    # IV. Thiết kế giao diện
    add_heading(doc, "IV. Thiết kế giao diện", level=2, size=14, space_before=12, space_after=6)
    add_paragraph(doc,
        "Giao diện được thiết kế theo phong cách hiện đại, sử dụng Tailwind CSS với bảng màu chủ đạo "
        "xanh dương (Blue-600: #2563EB) mang lại cảm giác chuyên nghiệp và đáng tin cậy. "
        "Font chữ sử dụng Inter của Google Fonts. Giao diện được tối ưu cho cả desktop và mobile "
        "(responsive với Tailwind breakpoints sm/md/lg)."
    )

    ui_screens = [
        ("Trang đăng nhập / đăng ký (Hình 3.9, 3.10)",
         "Thiết kế split-screen: bên trái là hình ảnh minh họa xe khách, bên phải là form. "
         "Form đăng nhập gồm: email, mật khẩu, nút 'Đăng nhập', link 'Chưa có tài khoản? Đăng ký ngay'. "
         "Form đăng ký gồm: email, mật khẩu, xác nhận mật khẩu, họ tên, SĐT. Validation inline hiển thị "
         "ngay khi nhập sai."),
        ("Trang tìm chuyến (Hình 3.11)",
         "Thiết kế card-based: thanh tìm kiếm nằm ngang với 3 input (điểm đi, điểm đến, ngày) + nút tìm. "
         "Kết quả hiển thị dạng danh sách card: thông tin tuyến, giờ khởi hành/đến, loại xe, giá, "
         "số ghế trống. Nút 'Chọn chuyến' ở mỗi card."),
        ("Trang chọn ghế (Hình 3.12)",
         "Layout 2 cột: bên trái là sơ đồ ghế dạng grid 5 cột, bên phải là bảng thông tin chuyến "
         "(tuyến, giờ, loại xe, giá vé). Ghế trống: xanh lá (#22c55e), đã đặt: đỏ (#ef4444), "
         "đang chọn: xanh dương (#3b82f6), ghế VIP: vàng. Bên dưới sơ đồ ghế là nút 'Tiếp tục'."),
        ("Trang thanh toán VNPay (Hình 3.13)",
         "Layout đơn giản: hiển thị thông tin vé, chọn phương thức thanh toán (VNPay / COD), "
         "nút 'Thanh toán'. Sau khi chọn VNPay, hiển thị loading spinner trong khi redirect. "
         "Trang kết quả hiển thị card vé với mã BUS-YYYYMMDD-XXXXX."),
        ("Admin Dashboard (Hình 3.14)",
         "Thiết kế sidebar + content: sidebar bên trái (menu: Dashboard, Users, Buses, Routes, Trips, "
         "Tickets, Assignments, Cargo, Maintenance). Content chính: 4 stat cards (Users, Buses, Routes, "
         "Trips today) + 2 biểu đồ (Recharts) + bảng cảnh báo bảo hiểm."),
        ("Trang quản lý Users (Hình 3.15)",
         "Layout bảng: cột: Username, Email, SĐT, Role, Status, Actions (Sửa, Khóa, Reset MK). "
         "Có nút 'Thêm người dùng' trên cùng. Tìm kiếm theo username/email."),
        ("Trang quản lý Buses (Hình 3.16)",
         "Tương tự Users: bảng cột: Biển số, Loại xe, Số ghế, Trạng thái, Đăng kiểm, Bảo hiểm, Actions. "
         "Có nút 'Thêm xe'. Màu status: xanh=ACTIVE, vàng=MAINTENANCE, đỏ=RETIRED."),
        ("Trang quản lý Routes (Hình 3.17)",
         "Bảng cột: ID, Điểm đi, Điểm đến, Khoảng cách, Thời gian, Giá vé, Actions."),
        ("Trang quản lý Trips (Hình 3.18)",
         "Bảng cột: ID, Tuyến, Xe, Giờ khởi hành, Giá, Trạng thái, Actions. Có filter theo ngày và trạng thái."),
        ("Trang phân công (Hình 3.19)",
         "Chọn chuyến từ danh sách, hiển thị thông tin chuyến, 2 dropdown: chọn tài xế + chọn phụ xe, "
         "nút 'Phân công'. Danh sách assignment hiển thị bên dưới."),
    ]
    for screen, desc in ui_screens:
        add_heading(doc, screen, level=3, size=13, bold=True, space_before=6, space_after=3)
        add_paragraph(doc, desc, space_before=2)

    # V. Thiết kế xử lý
    add_heading(doc, "V. Thiết kế xử lý", level=2, size=14, space_before=12, space_after=6)

    add_heading(doc, "1. Xử lý đặt vé (Ticket Booking)", level=3, size=13, bold=True, space_before=6, space_after=4)
    add_paragraph(doc,
        "Quy trình xử lý đặt vé sử dụng cơ chế Pessimistic Locking để đảm bảo tính nhất quán khi "
        "nhiều người cùng đặt 1 ghế. Khi nhận request POST /api/private/tickets, "
        "JwtAuthenticationFilter xác thực token → extract userId → chuyển đến TicketController. "
        "TicketService gọi seatRepository.findByIdForUpdate(seatId) — câu lệnh này sinh SQL "
        "'SELECT * FROM seats WHERE id=? FOR UPDATE', lock dòng trong database. "
        "Nếu ghế đã có Ticket active → ném BusinessConflictException. "
        "Nếu ghế trống → tạo Ticket(status=BOOKED), save. "
        "Nếu có exception DataIntegrityViolationException → ghế bị race condition → trả 409."
    )
    add_code(doc, "-- Pseudo-code xử lý đặt ghế\n"
                  "1. BEGIN TRANSACTION\n"
                  "2. SELECT * FROM seats WHERE id=? FOR UPDATE  -- Lock dòng\n"
                  "3. IF seat has active ticket THEN ROLLBACK + return 409\n"
                  "4. INSERT INTO tickets (trip_id, seat_id, passenger_id, status, ...)\n"
                  "5. COMMIT\n"
                  "6. RETURN ticket_response")

    add_heading(doc, "2. Xử lý thanh toán VNPay", level=3, size=13, bold=True, space_before=8, space_after=4)
    add_paragraph(doc,
        "Khi VNPay gửi IPN (Instant Payment Notification) đến /api/public/payment/vnpay/ipn, "
        "VnpayService thực hiện các bước: (1) Verify HMAC-SHA512 signature từ vnp_SecureHash; "
        "(2) Kiểm tra ResponseCode == '00'; (3) Kiểm tra idempotent: "
        "SELECT * FROM payments WHERE vnp_txn_ref=? AND status='SUCCESS' — nếu đã có → "
        "trả 200 OK ngay mà không xử lý lại; (4) Tìm Ticket tương ứng, cập nhật "
        "Payment→SUCCESS, Ticket→PAID; (5) Lưu thông tin bank code và response code; "
        "(6) Gửi SSE notification đến Admin Dashboard."
    )
    add_code(doc, "-- VNPay IPN Handler\n"
                  "1. Verify HMAC-SHA512(vnp_SecureHash)\n"
                  "2. IF NOT valid THEN return 'INVALID_SIGNATURE'\n"
                  "3. IF vnp_ResponseCode != '00' THEN return 'PAYMENT_FAILED'\n"
                  "4. IF payment already SUCCESS THEN return 'ALREADY_PROCESSED'  -- Idempotent\n"
                  "5. UPDATE payments SET status='SUCCESS', ... WHERE id=?\n"
                  "6. UPDATE tickets SET status='PAID' WHERE id=?\n"
                  "7. sseEmitter.send('PAYMENT_VNPAY_SUCCESS')\n"
                  "8. return 'OK'")

    add_heading(doc, "3. Xử lý phân công tài xế", level=3, size=13, bold=True, space_before=8, space_after=4)
    add_paragraph(doc,
        "Khi Admin phân công tài xế cho chuyến, TripAssignmentService kiểm tra: "
        "(1) Nhân viên đang ACTIVE; (2) Nhân viên chưa được phân công chuyến khác trùng thời gian. "
        "Câu truy vấn: SELECT * FROM trip_assignments ta "
        "JOIN trips t ON ta.trip_id=t.id "
        "WHERE ta.employee_id=? AND t.departure_time BETWEEN ? AND ? "
        "AND t.status IN ('SCHEDULED','RUNNING'). "
        "Nếu có kết quả → báo lỗi overlap. "
        "Nếu không → tạo TripAssignment."
    )

    add_heading(doc, "4. Xử lý SSE Notification", level=3, size=13, bold=True, space_before=8, space_after=4)
    add_paragraph(doc,
        "AdminNotificationService quản lý danh sách CopyOnWriteArrayList<SseEmitter>. "
        "Mỗi khi có sự kiện (booking mới, thanh toán thành công), service duyệt qua toàn bộ "
        "emitters và gửi event. CopyOnWriteArrayList đảm bảo thread-safety khi nhiều thread "
        "cùng thêm/bớt emitters. Khi client disconnect, emitter.complete() được gọi tự động "
        "hoặc qua try-catch IOException."
    )
    add_code(doc, "// AdminNotificationService\n"
                  "private List<SseEmitter> emitters = new CopyOnWriteArrayList<>();\n\n"
                  "public SseEmitter subscribe() {\n"
                  "  SseEmitter emitter = new SseEmitter(Long.MAX_VALUE);\n"
                  "  emitters.add(emitter);\n"
                  "  emitter.onCompletion(() -> emitters.remove(emitter));\n"
                  "  emitter.onTimeout(() -> emitters.remove(emitter));\n"
                  "  emitter.onError(e -> emitters.remove(emitter));\n"
                  "  return emitter;\n"
                  "}\n\n"
                  "public void notify(String event, Object data) {\n"
                  "  for (SseEmitter e : emitters) {\n"
                  "    try { e.send(SseEmitter.event().name(event).data(data)); }\n"
                  "    catch (IOException ex) { emitters.remove(e); }\n"
                  "  }\n"
                  "}")

    page_break(doc)

# ─────────────────────────────────────────────
# CHAPTER IV - PHÁT TRIỂN / THỰC THI
# ─────────────────────────────────────────────

def build_chapter4(doc):
    add_heading(doc, "CHƯƠNG IV. PHÁT TRIỂN / THỰC THI", level=1, size=16, space_before=6, space_after=12)

    # I. Màn hình đăng nhập / đăng ký
    add_heading(doc, "I. Màn hình đăng nhập / đăng ký", level=2, size=14, space_before=8, space_after=6)
    add_paragraph(doc,
        "Trang đăng nhập (Hình 4.1) là điểm khởi đầu của hệ thống. Giao diện chia làm 2 phần: "
        "bên trái là panel hình ảnh minh họa xe khách với nền gradient xanh dương, bên phải là "
        "form đăng nhập với logo 'XeKhách Pro' ở trên cùng. "
        "Form đăng nhập gồm: trường email với icon, trường mật khẩu có nút hiện/ẩn, "
        "nút 'Đăng nhập' màu xanh dương. Bên dưới có link 'Chưa có tài khoản? Đăng ký ngay'. "
        "Khi đăng nhập thành công, hệ thống lưu JWT token vào Zustand store (persist localStorage), "
        "redirect đến trang tìm chuyến (CUSTOMER) hoặc Admin Dashboard (ADMIN). "
        "Khi đăng nhập thất bại, hiển thị toast error: 'Email hoặc mật khẩu không đúng'. "
        "Trang đăng ký tương tự nhưng có thêm trường: họ tên, SĐT, xác nhận mật khẩu. "
        "Validation real-time: email format, SĐT 10 số, mật khẩu ≥ 8 ký tự, "
        "xác nhận mật khẩu phải khớp."
    )
    add_paragraph(doc,
        "Phía Backend, AuthController xử lý: POST /api/public/auth/login → JwtService tạo token "
        "chứa {sub: username, roles: [ROLE_ADMIN/ROLE_CUSTOMER], iat, exp}. "
        "POST /api/public/auth/register → UserService tạo User + Passenger, BCrypt hash password."
    )

    # II. Màn hình tìm chuyến & đặt vé
    add_heading(doc, "II. Màn hình tìm chuyến & đặt vé", level=2, size=14, space_before=10, space_after=6)
    add_paragraph(doc,
        "Trang tìm chuyến (Hình 4.2) là trang chính của khách hàng sau khi đăng nhập. "
        "Thanh tìm kiếm nằm ngang phía trên với 3 trường: điểm đi (select), điểm đến (select), "
        "ngày khởi hành (date picker). Nút 'Tìm kiếm' bên phải. "
        "Kết quả hiển thị dạng danh sách card: mỗi card chứa tuyến đường (Hà Nội → TP.HCM), "
        "giờ khởi hành, giờ đến (dựa trên estimatedDurationMin), loại xe (badge màu), "
        "số ghế trống (ví dụ 'còn 24 ghế'), giá vé (hiển thị nổi bật màu xanh), "
        "nút 'Chọn chuyến'. Card có hover effect để nổi bật."
    )
    add_paragraph(doc,
        "Sau khi chọn chuyến, hiển thị 3 bước tiếp theo liên tiếp trên cùng trang "
        "(Hình 4.3, 4.4, 4.5): "
        "(a) Chọn điểm đón: dropdown chọn thành phố → hiển thị danh sách điểm cụ thể bên dưới; "
        "(b) Chọn điểm trả: tương tự; "
        "(c) Chọn ghế: sơ đồ ghế grid 5 cột × n hàng, mỗi ghế là một ô vuông chứa số ghế. "
        "Ghế trống: nền xanh lá (#22c55e), chữ trắng. Ghế đã đặt: nền đỏ (#ef4444), chữ trắng, "
        "disabled. Ghế đang chọn: nền xanh dương (#3b82f6), chữ trắng. "
        "Ghế VIP: viền vàng (#eab308). Tổng giá hiển thị cố định ở dưới: 'Tổng: 500,000 VNĐ'. "
        "Khi chọn ghế → gọi API → tạo Ticket BOOKED → hiển thị form xác nhận. "
        "Form xác nhận: tóm tắt thông tin vé + 2 radio buttons: 'Thanh toán VNPay' / 'Thanh toán COD'."
    )

    # III. Màn hình thanh toán
    add_heading(doc, "III. Màn hình thanh toán", level=2, size=14, space_before=10, space_after=6)
    add_paragraph(doc,
        "Khi khách hàng chọn VNPay và nhấn 'Thanh toán' (Hình 4.6), hệ thống gọi "
        "POST /api/private/payment/vnpay/create. Backend tạo VNPay URL và trả về. "
        "Frontend redirect sang URL VNPay sandbox. Khách hàng chọn ngân hàng hoặc thanh toán "
        "bằng thẻ test (thẻ ATM: 9704195791369783 / OTP: 123456). "
        "Sau khi thanh toán, VNPay redirect về Return URL. Backend verify và cập nhật "
        "Payment + Ticket → PAID, gửi SSE notification đến Admin Dashboard. "
        "Frontend hiển thị trang kết quả: card vé màu xanh lá với mã vé BUS-20250627-00001, "
        "thông tin hành khách, chuyến xe, ghế, điểm đón/trả. Nếu thanh toán thất bại (mã ≠ 00), "
        "hiển thị card đỏ với thông báo lỗi và nút 'Thử lại'."
    )

    # IV. Admin Dashboard
    add_heading(doc, "IV. Màn hình Admin Dashboard", level=2, size=14, space_before=10, space_after=6)
    add_paragraph(doc,
        "Admin Dashboard (Hình 4.7) là trang tổng quan dành cho Quản trị viên sau khi đăng nhập. "
        "Giao diện gồm: Sidebar bên trái (menu điều hướng) + Content chính. "
        "4 Stat Cards ở trên: (1) Tổng người dùng (số + biểu tượng user); "
        "(2) Tổng xe khách (số + biểu tượng bus); "
        "(3) Tổng tuyến đường (số + biểu tượng map); "
        "(4) Chuyến hôm nay (số + biểu tượng calendar). "
        "2 Biểu đồ Recharts: Pie chart phân bổ loại xe (Limousine/Sleeper/Seat) và "
        "Bar chart số nhân viên theo loại (Driver/Assistant/Technician/Dispatcher/Manager). "
        "Bảng cảnh báo bảo hiểm: hiển thị các xe có insurance_expiry trong vòng 30 ngày, "
        "màu nền vàng nhạt. "
        "SSE notification: khi có booking mới hoặc thanh toán VNPay thành công, "
        "toast popup xuất hiện phía trên bên phải: "
        "'🔔 Đặt vé mới: Mã BUS-20250627-00001 - Khách Nguyễn Văn A'."
    )

    # V. Quản lý chuyến & phân công
    add_heading(doc, "V. Màn hình quản lý chuyến xe & phân công", level=2, size=14, space_before=10, space_after=6)
    add_paragraph(doc,
        "Trang quản lý chuyến xe (Hình 4.8) hiển thị bảng tất cả chuyến với các cột: "
        "ID, Tuyến (Hà Nội → TP.HCM), Xe (29B-12345), Giờ khởi hành, Giờ đến, Giá vé, "
        "Trạng thái (badge màu: xanh=SCHEDULED, cam=RUNNING, xám=COMPLETED, đỏ=CANCELLED), "
        "Actions (Sửa, Phân công, Hủy). "
        "Nút 'Thêm chuyến' mở modal: chọn tuyến, chọn xe (chỉ hiện xe ACTIVE), "
        "chọn giờ khởi hành. Backend kiểm tra overlap trước khi lưu. "
        "Trang phân công (Hình 4.9): chọn chuyến từ danh sách → hiển thị thông tin chuyến. "
        "Dropdown 'Tài xế' hiển thị danh sách nhân viên loại DRIVER (sắp xếp theo kinh nghiệm giảm dần). "
        "Dropdown 'Phụ xe' hiển thị nhân viên loại ASSISTANT. "
        "Nút 'Phân công' gọi API → kiểm tra overlap → tạo TripAssignment. "
        "Danh sách assignment hiển thị bên dưới: tên tài xế, vai trò, nút xóa."
    )

    page_break(doc)

# ─────────────────────────────────────────────
# CHAPTER V - TRIỂN KHAI
# ─────────────────────────────────────────────

def build_chapter5(doc):
    add_heading(doc, "CHƯƠNG V. TRIỂN KHAI", level=1, size=16, space_before=6, space_after=12)

    # I. Cài đặt
    add_heading(doc, "I. Cài đặt", level=2, size=14, space_before=8, space_after=6)

    add_heading(doc, "1. Môi trường yêu cầu", level=3, size=13, bold=True, space_before=6, space_after=4)
    env = [
        "Java Development Kit (JDK) phiên bản 17 trở lên.",
        "Node.js phiên bản 18 trở lên.",
        "MySQL Server phiên bản 8.0 trở lên.",
        "Trình duyệt web: Chrome / Firefox / Edge phiên bản mới nhất.",
        "Git (để clone source code).",
        "Maven phiên bản 3.9+ (để build backend).",
        "VNPay sandbox account (đăng ký tại sandbox.vnpayment.vn).",
    ]
    for e in env:
        add_bullet(doc, e, size=13)

    add_heading(doc, "2. Các bước cài đặt", level=3, size=13, bold=True, space_before=8, space_after=4)

    steps_install = [
        ("Bước 1: Clone source code",
         "Tải toàn bộ source code từ thư mục dự án. Cấu trúc gồm 2 thư mục: backend/ và frontend/."),
        ("Bước 2: Cài đặt MySQL",
         "Cài MySQL Server 8.0. Tạo database 'bus_management_db'. "
         "Cấu hình username='root', password='123456' trong backend/src/main/resources/application.properties. "
         "Database sẽ tự động được tạo bảng khi Spring Boot khởi động (JPA auto-ddl=update). "
         "DataInitializer sẽ tự động seed dữ liệu mẫu (users, routes, buses, employees, pickup/dropoff points)."),
        ("Bước 3: Cài đặt Backend",
         "Mở terminal tại thư mục backend/. Chạy lệnh: mvn clean install. "
         "Sau đó chạy: mvn spring-boot:run. Backend khởi động ở port 8080. "
         "Kiểm tra: GET http://localhost:8080/api/health → trả về 'OK'."),
        ("Bước 4: Cài đặt Frontend",
         "Mở terminal tại thư mục root (frontend/). Chạy: npm install. "
         "Sau đó: npm run dev. Frontend khởi động ở port 4173. "
         "Truy cập: http://localhost:4173."),
        ("Bước 5: Cấu hình VNPay",
         "Đăng ký tài khoản VNPay sandbox tại sandbox.vnpayment.vn. "
         "Lấy mã TMN (Terminal ID) và Hash Secret. "
         "Cập nhật vào application.properties: app.vnpay.tmn-code và app.vnpay.hash-secret. "
         "Return URL: http://localhost:4173/payment/return (cấu hình trên sandbox portal)."),
    ]
    for title, desc in steps_install:
        add_heading(doc, title, level=3, size=13, bold=True, space_before=4, space_after=2)
        add_paragraph(doc, desc, space_before=2)

    add_heading(doc, "3. Bảng cài đặt chức năng", level=3, size=13, bold=True, space_before=8, space_after=6)
    add_paragraph(doc, "Bảng 5.1: Tình trạng cài đặt các chức năng", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=3, space_after=4)
    install_data = [
        ["STT", "Chức năng", "Mức độ hoàn thành", "Ghi chú"],
        ["1", "Đăng nhập / Đăng ký", "100%", "JWT auth, BCrypt password"],
        ["2", "Tra cứu & tìm kiếm chuyến xe", "100%", "Lọc theo điểm đi, đến, ngày"],
        ["3", "Xem sơ đồ ghế & trạng thái ghế", "100%", "Real-time từ DB"],
        ["4", "Đặt vé xe khách", "100%", "Race-condition-safe (SELECT FOR UPDATE)"],
        ["5", "Thanh toán VNPay", "100%", "Sandbox test, HMAC-SHA512"],
        ["6", "Thanh toán COD", "100%", "Xác nhận tức thì, SSE notification"],
        ["7", "Xem vé của tôi", "100%", "Lọc theo customer"],
        ["8", "Hủy vé", "100%", "Kiểm tra thời hạn"],
        ["9", "Quản lý Dashboard Admin", "100%", "Stats, biểu đồ Recharts, SSE"],
        ["10", "CRUD Users", "100%", "Khóa/mở khóa, reset mật khẩu"],
        ["11", "CRUD Buses", "100%", "Auto tạo seats, cảnh báo bảo hiểm"],
        ["12", "CRUD Routes", "100%", "Auto tính giá"],
        ["13", "CRUD Trips", "100%", "Overlap detection"],
        ["14", "Phân công tài xế & phụ xe", "100%", "Overlap detection assignment"],
        ["15", "CRUD Employees", "100%", "Lọc theo loại"],
        ["16", "Quản lý vé (Admin)", "100%", "Xác nhận COD, hủy vé"],
        ["17", "CRUD Maintenance", "100%", "Ghi nhận bảo trì xe"],
        ["18", "CRUD Cargo", "100%", "Quản lý hàng hóa"],
        ["19", "SSE Real-time Notification", "100%", "Booking + VNPay success"],
        ["20", "Audit Log", "100%", "Ghi nhật ký thao tác Admin"],
        ["21", "Cập nhật hồ sơ", "100%", "Customer profile"],
    ]
    t = add_table_bordered(doc, len(install_data), 4, install_data)
    for cell in t.rows[0].cells:
        shade_cell(cell, "D9E1F2")
        cell.paragraphs[0].runs[0].bold = True

    # II. Thử nghiệm
    add_heading(doc, "II. Thử nghiệm", level=2, size=14, space_before=12, space_after=6)

    add_heading(doc, "1. Tài khoản test", level=3, size=13, bold=True, space_before=6, space_after=6)
    add_paragraph(doc, "Bảng 5.2: Tài khoản dùng để test", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=3, space_after=4)
    test_data = [
        ["Vai trò", "Username", "Mật khẩu", "Ghi chú"],
        ["Quản trị viên (ADMIN)", "admin", "ChangeMe@123",
         "Full quyền: CRUD users, buses, routes, trips, assignments, tickets, employees, cargo, maintenance, xem dashboard, nhận SSE notification."],
        ["Khách hàng (CUSTOMER)", "customer1", "password123",
         "Đặt vé, thanh toán VNPay/COD, xem vé của tôi, cập nhật hồ sơ."],
        ["Khách hàng (CUSTOMER)", "customer2", "password123",
         "Tương tự customer1, dùng để test đặt nhiều vé."],
        ["Khách hàng (CUSTOMER)", "customer3", "password123",
         "Test đặt vé và hủy vé."],
    ]
    t = add_table_bordered(doc, len(test_data), 4, test_data)
    for cell in t.rows[0].cells:
        shade_cell(cell, "D9E1F2")
        cell.paragraphs[0].runs[0].bold = True

    add_heading(doc, "2. Kịch bản test chi tiết", level=3, size=13, bold=True, space_before=8, space_after=6)

    test_cases = [
        ("TC01 - Đặt vé + thanh toán VNPay",
         "Đăng nhập customer1 → Tìm chuyến (Hà Nội → TP.HCM, ngày hiện tại) → "
         "Chọn ghế A1 → Chọn điểm đón (Hà Nội - Bến xe Mỹ Đình) → Chọn điểm trả (TP.HCM - Bến xe Miền Đông) → "
         "Chọn thanh toán VNPay → Nhấn 'Thanh toán' → Redirect sang VNPay sandbox → "
         "Chọn ngân hàng (Vietcombank) → Nhập thẻ test (9704195791369783) → OTP (123456) → "
         "Xác nhận → Redirect về Return URL → Vé trạng thái PAID. "
         "Kiểm tra: Admin Dashboard nhận SSE notification 'Thanh toán VNPay thành công'."),
        ("TC02 - Đặt vé + thanh toán COD",
         "Đăng nhập customer2 → Tìm chuyến → Chọn ghế A2 → Chọn điểm đón/trả → "
         "Chọn thanh toán COD → Xác nhận → Vé trạng thái CONFIRMED. "
         "Kiểm tra: Admin Dashboard nhận SSE notification 'Đặt vé mới'."),
        ("TC03 - Hủy vé",
         "Đăng nhập customer2 → Vào 'Vé của tôi' → Tìm vé CONFIRMED → Nhấn 'Hủy vé' → "
         "Xác nhận hủy → Vé chuyển sang CANCELLED. "
         "Kiểm tra: ghế A2 trở về trạng thái trống trong sơ đồ ghế."),
        ("TC04 - Test race condition (đặt cùng 1 ghế)",
         "Mở 2 trình duyệt (2 tài khoản customer2 và customer3) → Cùng chọn ghế A3 → "
         "Cả 2 đều nhấn 'Đặt' gần như cùng lúc → "
         "1 trong 2 nhận thông báo 'Ghế đã được đặt, vui lòng chọn ghế khác' (HTTP 409). "
         "1 trong 2 đặt thành công → ghế A3 chuyển đỏ trên trình duyệt còn lại."),
        ("TC05 - Quản lý chuyến & phân công tài xế",
         "Đăng nhập admin → Quản lý chuyến xe → 'Thêm chuyến' → Chọn tuyến Hà Nội→TP.HCM, "
         "xe 29B-12345, giờ khởi hành ngày mai 07:00 → Lưu → "
         "Vào 'Phân công' → Chọn chuyến → Gán tài xế Lê Văn Minh, phụ xe Trần Văn Hùng → 'Phân công'. "
         "Kiểm tra: danh sách assignment hiển thị đúng tên, vai trò."),
        ("TC06 - Cảnh báo bảo hiểm",
         "Đăng nhập admin → Dashboard → Bảng 'Cảnh báo bảo hiểm' hiển thị các xe có "
         "insurance_expiry trong vòng 30 ngày với nền vàng. "
         "Test: thêm 1 xe mới với insurance_expiry = ngày hôm nay + 7 ngày → "
         "Dashboard tự động hiển thị cảnh báo mà không cần reload."),
    ]
    for title, desc in test_cases:
        add_heading(doc, title, level=3, size=13, bold=True, space_before=6, space_after=3)
        add_paragraph(doc, desc, space_before=2)

    page_break(doc)

# ─────────────────────────────────────────────
# CHAPTER VI - KẾT LUẬN
# ─────────────────────────────────────────────

def build_chapter6(doc):
    add_heading(doc, "CHƯƠNG VI. KẾT LUẬN", level=1, size=16, space_before=6, space_after=12)

    # I. Kết quả
    add_heading(doc, "I. Kết quả đã thực hiện", level=2, size=14, space_before=8, space_after=6)
    add_paragraph(doc,
        "Sau thời gian nghiên cứu, phân tích, thiết kế và triển khai, hệ thống Bus Management "
        "(XeKhách Pro) đã hoàn thành các mục tiêu đề ra. Cụ thể:"
    )

    results = [
        "Hoàn thành xây dựng hệ thống Backend Spring Boot 3.2 với 11 REST Controllers, 9 Services, "
        "14 Repositories, 15 JPA Entities, 31 DTOs, kiến trúc MVC phân tách rõ ràng Controller → Service → Repository.",
        "Hoàn thành giao diện Frontend React 18 với 16 trang web, thiết kế responsive, trải nghiệm "
        "người dùng mượt mà, sử dụng Tailwind CSS và Zustand state management.",
        "Tích hợp thành công cổng thanh toán VNPay Sandbox với mã hóa HMAC-SHA512, xử lý "
        "idempotent và server-to-server IPN callback.",
        "Triển khai thành công cơ chế race-condition-safe booking với Pessimistic Locking "
        "(SELECT FOR UPDATE) đảm bảo không có 2 khách nào đặt cùng 1 ghế.",
        "Xây dựng Admin Dashboard với thống kê realtime, biểu đồ Recharts, "
        "và hệ thống thông báo SSE real-time khi có đặt vé hoặc thanh toán.",
        "Hoàn thiện bộ công cụ quản lý cho Admin: quản lý users, buses, routes, trips, "
        "assignments, employees, tickets, maintenance, cargo, audit log.",
        "Cài đặt đầy đủ cơ chế bảo mật: JWT authentication, BCrypt password hashing, "
        "Spring Security authorization, CORS configuration.",
        "Khởi tạo dữ liệu mẫu phong phú: 20 xe bus (5 loại), 4 tuyến đường, "
        "10 tài xế + 10 phụ xe, 80+ điểm đón/trả, 5 tài khoản test.",
        "Viết đầy đủ báo cáo đồ án CNPM theo đúng format mẫu của Học viện.",
    ]
    for r in results:
        add_bullet(doc, r, size=13)

    # II. Ưu khuyết điểm
    add_heading(doc, "II. Ưu khuyết điểm", level=2, size=14, space_before=12, space_after=6)

    add_heading(doc, "1. Ưu điểm", level=3, size=13, bold=True, space_before=6, space_after=4)
    pros = [
        "Giao diện hiện đại, thân thiện với người dùng, responsive trên mọi thiết bị.",
        "Kiến trúc phân lớp rõ ràng, code có tổ chức tốt, dễ bảo trì và mở rộng.",
        "Đảm bảo tính nhất quán dữ liệu (concurrency control) trong quá trình đặt vé.",
        "Tích hợp thanh toán trực tuyến đầy đủ (VNPay) và hỗ trợ COD linh hoạt.",
        "Thông báo real-time giúp Admin nắm bắt sự kiện kịp thời.",
        "Audit log ghi nhận mọi thao tác, đảm bảo tính truy vết.",
        "Có dữ liệu mẫu phong phú, thuận tiện cho việc kiểm thử.",
    ]
    for p in pros:
        add_bullet(doc, p, size=13)

    add_heading(doc, "2. Khuyết điểm", level=3, size=13, bold=True, space_before=8, space_after=4)
    cons = [
        "Chưa triển khai ứng dụng mobile (Android/iOS), chỉ có giao diện web.",
        "Chưa tích hợp gửi email xác nhận đặt vé tự động cho hành khách.",
        "Chưa triển khai thông báo SMS khi có cập nhật trạng thái vé.",
        "Chưa hỗ trợ đa ngôn ngữ (i18n) cho giao diện.",
        "Chưa triển khai in vé điện tử (PDF ticket) để khách hàng tải về.",
        "VNPay chỉ ở chế độ Sandbox, chưa triển khai production với certificate thật.",
        "Chưa triển khai CI/CD pipeline (GitHub Actions, Docker) cho việc deploy.",
        "Chưa có module báo cáo thống kê doanh thu chi tiết theo tháng/quý.",
    ]
    for c in cons:
        add_bullet(doc, c, size=13)

    # III. Hướng mở rộng
    add_heading(doc, "III. Hướng mở rộng trong tương lai", level=2, size=14, space_before=12, space_after=6)
    extensions = [
        "Phát triển ứng dụng di động Android/iOS sử dụng React Native hoặc Flutter, "
        "giúp hành khách đặt vé mọi lúc mọi nơi từ smartphone.",
        "Tích hợp dịch vụ gửi email tự động qua SMTP (JavaMailSender) hoặc SendGrid: "
        "xác nhận đặt vé, nhắc lịch trình trước 24 giờ, thông báo hủy/chậm chuyến.",
        "Tích hợp SMS gateway (Viettel, VNPT) để gửi thông báo qua tin nhắn SMS.",
        "Triển khai in vé PDF với mã QR code, khách hàng có thể tải vé về hoặc "
        "scan QR tại quầy làm thủ tục.",
        "Tích hợp Google Maps API để hiển thị vị trí xe realtime trên bản đồ "
        "(cần trang bị GPS tracker trên xe).",
        "Bổ sung module thống kê doanh thu: biểu đồ doanh thu theo tháng/quý/năm, "
        "báo cáo tỷ lệ lấp đầy ghế (occupancy rate), top tuyến phổ biến.",
        "Triển khai Docker containerization và CI/CD pipeline với GitHub Actions "
        "để tự động build, test và deploy lên server cloud (AWS, Azure, hoặc VNG Cloud).",
        "Tích hợp Single Sign-On (SSO) với tài khoản Google/Facebook để đăng nhập nhanh hơn.",
        "Xây dựng module khuyến mãi: mã giảm giá, voucher, chương trình tích điểm cho khách thân thiện.",
        "Nâng cao bảo mật: two-factor authentication (2FA), rate limiting, "
        "security audit log với alerting tự động.",
    ]
    for e in extensions:
        add_bullet(doc, e, size=13)

    page_break(doc)

# ─────────────────────────────────────────────
# REFERENCES
# ─────────────────────────────────────────────

def build_references(doc):
    add_heading(doc, "TÀI LIỆU THAM KHẢO", level=1, size=16, space_before=6, space_after=12)

    refs = [
        "[1] Spring Boot Documentation. Truy cập: https://spring.io/projects/spring-boot",
        "[2] Spring Security Reference. Truy cập: https://docs.spring.io/spring-security/reference/",
        "[3] React Documentation. Truy cập: https://react.dev/",
        "[4] Vite Guide. Truy cập: https://vitejs.dev/guide/",
        "[5] VNPay Payment Gateway API Documentation v2.1.0. Truy cập: https://sandbox.vnpayment.vn/apis/",
        "[6] JWT (JSON Web Token) - RFC 7519. Truy cập: https://datatracker.ietf.org/rfc/rfc7519/",
        "[7] MySQL 8.0 Reference Manual. Truy cập: https://dev.mysql.com/doc/refman/8.0/en/",
        "[8] Tailwind CSS Documentation. Truy cập: https://tailwindcss.com/docs/",
        "[9] Hibernate ORM Documentation. Truy cập: https://docs.jboss.org/hibernate/orm/current/",
        "[10] Server-Sent Events (SSE) - MDN Web Docs. Truy cập: https://developer.mozilla.org/en-US/docs/Web/API/Server-sent_events",
        "[11] Recharts Documentation. Truy cập: https://recharts.org/en-US/",
        "[12] Zustand State Management. Truy cập: https://zustand.docs.pmnd.rs/",
        "[13] Lombok Project. Truy cập: https://projectlombok.org/",
        "[14] Java Persistence API (JPA) 3.1 Specification. Truy cập: https://jakarta.ee/specifications/persistence/3.1/",
    ]
    for ref in refs:
        add_paragraph(doc, ref, space_before=4, space_after=4)

# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────

def main():
    doc = Document()

    # Page setup
    for section in doc.sections:
        set_page_setup(section)

    # Build all chapters
    build_cover(doc)
    build_toc(doc)
    build_figures_tables(doc)
    build_abbreviations(doc)
    build_chapter1(doc)
    build_chapter2(doc)
    build_chapter3(doc)
    build_chapter4(doc)
    build_chapter5(doc)
    build_chapter6(doc)
    build_references(doc)

    # Save
    out_path = r'd:\metbus\BUS\Bao Cao Do An CNPM - XeKhach Pro.docx'
    doc.save(out_path)
    print(f"✅ Saved: {out_path}")

if __name__ == '__main__':
    main()
